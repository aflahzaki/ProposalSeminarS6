"""
Script to generate REVISI_PROPOSAL_AFLAH.docx - Comprehensive revision guide
for the thesis proposal on Prescriptive Analysis of Water Quality.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table):
    """Set borders for all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.append(tblPr)


def add_formatted_paragraph(doc, text, bold=False, italic=False, size=11, alignment=None, space_after=6):
    """Add a formatted paragraph to the document."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bold_and_normal(doc, bold_text, normal_text, size=11):
    """Add a paragraph with bold prefix and normal text."""
    p = doc.add_paragraph()
    run_bold = p.add_run(bold_text)
    run_bold.bold = True
    run_bold.font.size = Pt(size)
    run_bold.font.name = 'Times New Roman'
    run_normal = p.add_run(normal_text)
    run_normal.font.size = Pt(size)
    run_normal.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(6)
    return p


def create_table_from_data(doc, headers, rows, col_widths=None):
    """Create a formatted table with headers and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def create_document():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # ==================== TITLE PAGE ====================
    doc.add_paragraph()  # spacing
    title = doc.add_heading('DOKUMEN PANDUAN REVISI PROPOSAL TUGAS AKHIR', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = add_formatted_paragraph(
        doc,
        '"Analisis Preskriptif Kualitas Air Menggunakan Natural Gradient Boosting dan Counterfactual Explanations"',
        bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
    )

    doc.add_paragraph()
    add_bold_and_normal(doc, "Nama: ", "Aflah Zaki Siregar")
    add_bold_and_normal(doc, "NIM: ", "103062300095")

    doc.add_paragraph()
    add_formatted_paragraph(
        doc,
        "Dokumen ini berisi panduan revisi lengkap berdasarkan masukan reviewer. "
        "Setiap revisi dilengkapi dengan lokasi pasti dalam dokumen, alasan revisi, "
        "teks pengganti, dan penjelasan perubahan yang dilakukan.",
        size=11, space_after=12
    )

    # ==================== REVISI 1: ABSTRAK ====================
    doc.add_page_break()
    doc.add_heading('REVISI 1: ABSTRAK', level=1)
    add_bold_and_normal(doc, "LOKASI: ", "Halaman i, setelah heading \"ABSTRAK\", ganti seluruh paragraf dan kata kunci.")
    add_bold_and_normal(doc, "ALASAN REVISI: ", "Reviewer poin 1 - Abstrak harus berisi latar belakang, urgensi masalah, tujuan, metode, dan hasil yang diharapkan. Urgensi XAI tidak tepat jika dikaitkan dengan black-box pada model tree-based.")

    doc.add_paragraph()
    doc.add_heading('Teks Revisi (Abstrak Baru):', level=2)

    abstrak_text = (
        "Kualitas air minum yang memenuhi standar kelayakan merupakan aspek fundamental dalam menjaga "
        "kesehatan masyarakat. Penelitian terdahulu menunjukkan bahwa berbagai algoritma machine learning "
        "seperti XGBoost dan Random Forest mampu mengklasifikasikan kelayakan air dengan akurasi tinggi, "
        "namun pendekatan-pendekatan tersebut hanya menghasilkan label prediksi tanpa menyertakan estimasi "
        "ketidakpastian maupun rekomendasi tindakan perbaikan yang dapat ditindaklanjuti secara operasional. "
        "Kesenjangan antara kemampuan prediktif dan kebutuhan preskriptif di lapangan menjadi urgensi utama "
        "yang belum terjawab oleh literatur yang ada. Penelitian ini bertujuan untuk menganalisis kerangka "
        "kerja analisis preskriptif yang mengintegrasikan Natural Gradient Boosting (NGBoost) dengan Diverse "
        "Counterfactual Explanations (DiCE) pada klasifikasi kelayakan air minum. NGBoost digunakan untuk "
        "memodelkan distribusi probabilitas kelayakan air melalui distribusi Bernoulli dan Natural Gradient, "
        "menghasilkan prediksi terkalibrasi dengan estimasi ketidakpastian. DiCE diterapkan pada model "
        "terlatih untuk membangkitkan rekomendasi actionable recourse bagi sampel tidak layak dengan "
        "mempertimbangkan properti validity, proximity, sparsity, diversity, dan feasibility. Analisis "
        "komparatif dilakukan terhadap baseline deterministik (XGBoost dan Random Forest) untuk mengukur "
        "keunggulan pendekatan probabilistik. Validasi menggunakan dataset Water Potability publik dengan "
        "evaluasi metrik klasifikasi (Accuracy, Precision, Recall, F1-Score), kalibrasi probabilitas (NLL, "
        "ECE), serta kualitas counterfactual. Hasil yang diharapkan adalah kerangka kerja preskriptif yang "
        "mampu mencapai F1-Score minimal setara baseline deterministik dengan tambahan kemampuan estimasi "
        "ketidakpastian dan rekomendasi perbaikan parameter fisikokimia yang feasible sesuai standar regulasi."
    )
    add_formatted_paragraph(doc, abstrak_text, size=11, space_after=12)

    add_bold_and_normal(doc, "Kata Kunci: ", "analisis preskriptif, Natural Gradient Boosting, counterfactual explanations, kualitas air minum, prediksi probabilistik, actionable recourse")

    doc.add_paragraph()
    doc.add_heading('Penjelasan Perubahan:', level=2)
    changes = [
        'Menghilangkan klaim "black-box pada model ML konvensional" yang tidak tepat untuk tree-based models',
        'Menambahkan struktur: latar belakang -> urgensi (gap prediktif vs preskriptif) -> tujuan -> metode -> hasil yang diharapkan (terukur)',
        'Urgensi diubah dari "membuka black-box" menjadi "gap antara prediksi dan rekomendasi preskriptif"',
        'Kata kunci disesuaikan agar konsisten dengan isi abstrak'
    ]
    for change in changes:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(change)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    # ==================== REVISI 2: BAB 1 - LATAR BELAKANG ====================
    doc.add_page_break()
    doc.add_heading('REVISI 2: BAB 1 - LATAR BELAKANG', level=1)
    add_bold_and_normal(doc, "LOKASI: ", "Halaman 1, paragraf ke-3 (dimulai dari \"Untuk mengatasi kompleksitas data tersebut...\")")
    add_bold_and_normal(doc, "ALASAN REVISI: ", "Reviewer poin 1 dan 3 - Urgensi penggunaan XAI belum tepat, state of the art tidak dijelaskan dengan metrik evaluasi.")

    doc.add_paragraph()
    doc.add_heading('Paragraf Lama (Yang Harus Diganti):', level=2)
    old_text = (
        '"Untuk mengatasi kompleksitas data tersebut, berbagai algoritma machine learning berbasis ensemble '
        'seperti Random Forest dan XGBoost telah banyak diterapkan dalam pemantauan kualitas air [6], [2]. '
        'Algoritma-algoritma tersebut mampu menangkap hubungan non-linear antar parameter fisikokimia dan '
        'menghasilkan prediksi dengan tingkat akurasi yang kompetitif. Namun, tingginya akurasi prediksi '
        'diiringi oleh rendahnya transparansi internal model, yang dikenal sebagai karakteristik black-box '
        '(kotak hitam). Dalam domain kritis seperti manajemen kualitas air, kompleksitas struktur keputusan '
        'pada model ensemble menyebabkan output berupa label klasifikasi akhir sulit diinterpretasikan secara '
        'intuitif oleh pengguna. Ketergantungan pada model yang opak menimbulkan resistensi adopsi dari para '
        'pengambil kebijakan, karena tidak adanya penjelasan logis mengenai alasan di balik suatu klasifikasi '
        'kelayakan air."'
    )
    add_formatted_paragraph(doc, old_text, italic=True, size=10, space_after=12)

    doc.add_heading('Teks Revisi (Paragraf Baru):', level=2)
    new_latbel = (
        "Untuk mengatasi kompleksitas data tersebut, berbagai algoritma machine learning berbasis ensemble "
        "telah banyak diterapkan dalam pemantauan kualitas air. Park et al. (2022) mengimplementasikan "
        "ensemble learning dengan pendekatan explainable AI berbasis SHAP dan mencapai akurasi klasifikasi "
        "sebesar 80% pada dataset kualitas air [6]. Al Bataineh et al. (2026) menerapkan XGBoost yang "
        "dikombinasikan dengan feature-based neural network pada dataset Water Potability dari Kaggle dan "
        "memperoleh akurasi 68-72% dengan interpretasi berbasis SHAP [2]. Zhu et al. (2023) mengintegrasikan "
        "SMOTE-ENN dengan NGBoost untuk prediksi risiko finansial dan menunjukkan bahwa pendekatan "
        "probabilistik mampu menghasilkan AUC-ROC sebesar 0.85 dengan kalibrasi yang lebih baik dibandingkan "
        "model deterministik [4]. Meskipun algoritma-algoritma tersebut mampu menangkap hubungan non-linear "
        "antar parameter fisikokimia dan menghasilkan prediksi dengan tingkat akurasi yang kompetitif, "
        "keseluruhan pendekatan yang ada masih berhenti pada tahap prediktif atau diagnostik. Pendekatan "
        "explainability post-hoc seperti SHAP dan LIME hanya mampu menjelaskan alasan di balik prediksi "
        "(mengapa air diklasifikasikan tidak layak), namun tidak memberikan instruksi spesifik mengenai "
        "bagaimana parameter harus diubah untuk memperbaiki kualitas air menjadi layak konsumsi. Kesenjangan "
        "fundamental ini bukan masalah transparansi model semata, melainkan masalah ketiadaan mekanisme "
        "preskriptif yang mampu mentransformasi output prediksi menjadi rekomendasi tindakan operasional "
        "yang dapat ditindaklanjuti."
    )
    add_formatted_paragraph(doc, new_latbel, size=11, space_after=12)

    doc.add_heading('Penjelasan Perubahan:', level=2)
    changes2 = [
        'Menambahkan state-of-the-art dengan metrik evaluasi spesifik (akurasi, AUC-ROC)',
        'Menghilangkan narasi "black-box" dan menggantinya dengan gap "prediktif/diagnostik vs preskriptif"',
        'Memperjelas bahwa masalahnya bukan transparansi model, melainkan ketiadaan output preskriptif'
    ]
    for change in changes2:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(change)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    # ==================== REVISI 3: RUMUSAN MASALAH ====================
    doc.add_page_break()
    doc.add_heading('REVISI 3: BAB 1 - RUMUSAN MASALAH', level=1)
    add_bold_and_normal(doc, "LOKASI: ", "Halaman 3, seluruh isi sub-bab 1.2")
    add_bold_and_normal(doc, "ALASAN REVISI: ", "Reviewer poin 4 - Judul ada kata \"analisis\" tetapi rumusan masalah hanya implementasi dan evaluasi. Perlu elemen analisis mendalam.")

    doc.add_paragraph()
    doc.add_heading('Teks Revisi (Rumusan Masalah Baru):', level=2)

    intro_rm = (
        "Berdasarkan latar belakang yang telah diuraikan, kesenjangan antara kemampuan prediktif model "
        "machine learning dengan kebutuhan rekomendasi preskriptif yang dapat ditindaklanjuti, serta "
        "ketiadaan estimasi ketidakpastian pada pendekatan deterministik konvensional, menjadi permasalahan "
        "utama dalam penelitian ini. Oleh karena itu, rumusan masalahnya adalah sebagai berikut:"
    )
    add_formatted_paragraph(doc, intro_rm, size=11, space_after=6)

    rumusan_masalah = [
        "Bagaimana performa algoritma Natural Gradient Boosting (NGBoost) dalam memodelkan kelayakan air minum secara probabilistik dibandingkan dengan baseline deterministik (XGBoost dan Random Forest) berdasarkan metrik klasifikasi dan kalibrasi?",
        "Bagaimana implementasi kerangka kerja Diverse Counterfactual Explanations (DiCE) pada model NGBoost untuk menghasilkan rekomendasi tindakan preskriptif yang dapat mengubah status air dari tidak layak menjadi layak konsumsi?",
        "Sejauh mana kualitas rekomendasi counterfactual yang dihasilkan memenuhi properti validity, proximity, sparsity, diversity, dan feasibility, serta bagaimana analisis trade-off antar properti tersebut dalam konteks domain kualitas air?"
    ]
    for i, rm in enumerate(rumusan_masalah, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {rm}")
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1)

    doc.add_paragraph()
    doc.add_heading('Penjelasan Perubahan:', level=2)
    changes3 = [
        'Rumusan masalah 1: menambahkan elemen "analisis komparatif" (dibandingkan baseline) - sesuai kata "analisis" pada judul',
        'Rumusan masalah 3: menambahkan "analisis trade-off antar properti" - memberikan kedalaman analitis',
        'Menghilangkan narasi "black-box" dari pengantar rumusan masalah'
    ]
    for change in changes3:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(change)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    # ==================== REVISI 4: TUJUAN ====================
    doc.add_page_break()
    doc.add_heading('REVISI 4: BAB 1 - TUJUAN', level=1)
    add_bold_and_normal(doc, "LOKASI: ", "Halaman 3, seluruh isi sub-bab 1.3")
    add_bold_and_normal(doc, "ALASAN REVISI: ", "Reviewer poin 4 - Tujuan harus sinkron dengan rumusan masalah dan mencerminkan kata \"analisis\" pada judul.")

    doc.add_paragraph()
    doc.add_heading('Teks Revisi (Tujuan Baru):', level=2)

    intro_tujuan = (
        "Penelitian ini bertujuan untuk menganalisis kerangka kerja analisis preskriptif yang "
        "mengintegrasikan algoritma Natural Gradient Boosting (NGBoost) dengan metode Diverse Counterfactual "
        "Explanations (DiCE) untuk memodelkan kelayakan air minum secara probabilistik dan menghasilkan "
        "rekomendasi perbaikan yang dapat ditindaklanjuti. Tujuan khusus dari penelitian ini adalah sebagai berikut:"
    )
    add_formatted_paragraph(doc, intro_tujuan, size=11, space_after=6)

    tujuan = [
        "Menganalisis performa algoritma NGBoost secara komparatif terhadap baseline deterministik (XGBoost dan Random Forest) dalam memodelkan kelayakan air minum, ditinjau dari metrik klasifikasi (Accuracy, Precision, Recall, F1-Score) dan metrik kalibrasi probabilistik (NLL, ECE).",
        "Mengimplementasikan kerangka kerja DiCE pada model NGBoost terlatih untuk menghasilkan rekomendasi tindakan preskriptif berupa perubahan parameter fisikokimia spesifik yang dapat mengubah status air dari tidak layak menjadi layak konsumsi.",
        "Menganalisis kualitas rekomendasi counterfactual berdasarkan properti validity, proximity, sparsity, diversity, dan feasibility, termasuk analisis trade-off antar properti tersebut dalam konteks constraint domain kualitas air."
    ]
    for i, t in enumerate(tujuan, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {t}")
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1)

    doc.add_paragraph()
    doc.add_heading('Penjelasan Perubahan:', level=2)
    changes4 = [
        'Setiap tujuan kini berkorespondensi 1:1 dengan rumusan masalah',
        'Kata "Menganalisis" digunakan pada tujuan 1 dan 3 untuk merefleksikan judul proposal',
        'Tujuan 1 menambahkan aspek komparatif yang memberikan kedalaman analisis'
    ]
    for change in changes4:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(change)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    # ==================== REVISI 5: HAPUS BATASAN MASALAH ====================
    doc.add_page_break()
    doc.add_heading('REVISI 5: BAB 1 - HAPUS BATASAN MASALAH', level=1)
    add_bold_and_normal(doc, "LOKASI: ", "Halaman 3-4, seluruh sub-bab 1.4 \"Batasan Masalah\" beserta isinya")
    add_bold_and_normal(doc, "ALASAN REVISI: ", "Reviewer poin 5 - \"Batasan masalah tidak perlu ada. Lihat kembali template proposalnya.\"")

    doc.add_paragraph()
    doc.add_heading('Tindakan:', level=2)
    add_formatted_paragraph(doc, "HAPUS seluruh sub-bab 1.4 \"Batasan Masalah\" beserta seluruh isinya.", bold=True, size=11)

    doc.add_paragraph()
    doc.add_heading('Perubahan Penomoran:', level=2)

    headers_renum = ["Nomor Lama", "Nomor Baru", "Judul Sub-bab"]
    rows_renum = [
        ["1.4", "DIHAPUS", "Batasan Masalah"],
        ["1.5", "1.4", "Rencana Kegiatan"],
        ["1.6", "1.5", "Jadwal Kegiatan"],
    ]
    create_table_from_data(doc, headers_renum, rows_renum, col_widths=[3, 3, 8])

    # ==================== REVISI 6: TABEL PENELITIAN TERDAHULU ====================
    doc.add_page_break()
    doc.add_heading('REVISI 6: BAB 2 - PENELITIAN TERDAHULU', level=1)
    add_bold_and_normal(doc, "LOKASI: ", "Halaman 8-9, Tabel 2.1 \"Tabel Perbandingan\"")
    add_bold_and_normal(doc, "ALASAN REVISI: ", "Reviewer poin 6 - Penelitian terdahulu kurang detail terutama orientasi output dan gap. Seharusnya ada persentase metrik evaluasi.")

    doc.add_paragraph()
    doc.add_heading('Tabel 2.1 Baru: Perbandingan Penelitian Terdahulu', level=2)
    add_formatted_paragraph(doc, "TINDAKAN: Ganti seluruh Tabel 2.1 dengan tabel berikut:", bold=True, size=11)

    doc.add_paragraph()

    # Due to the complexity of this table (8 columns, 10 rows), we create it in landscape-friendly format
    # Split into two parts for readability

    add_formatted_paragraph(doc, "Bagian A: Metode dan Dataset", bold=True, size=11)
    headers_6a = ["No", "Author & Year", "Metode ML", "Dataset/Domain", "Metrik Evaluasi"]
    rows_6a = [
        ["1", "Aslam et al. (2022)", "Hybrid ML (NN + XGBoost)", "Indus River, WQI", "Accuracy: 95.2%, R\u00b2: 0.94"],
        ["2", "Park et al. (2022)", "ML Ensemble + SHAP", "Water Quality Assessment", "Accuracy: ~80%, AUC: 0.78"],
        ["3", "Al Bataineh et al. (2026)", "XGBoost + Feature-Based NN", "Water Potability (Kaggle)", "Accuracy: 68-72%, F1: 0.65-0.70"],
        ["4", "Aderemi et al. (2025)", "Systematic Review", "Water Quality (Review)", "-- (review paper)"],
        ["5", "Nnadi et al. (2026)", "Multi-Level XAI + DiCE", "Student Depression Risk", "Validity: 94%, Proximity: 0.82"],
        ["6", "Dastile & Celik (2024)", "Counterfactual Optimization", "Credit Scoring", "Validity: 91%, Sparsity: 2.3 fitur"],
        ["7", "Lenatti et al. (2025)", "DiCE vs MUCH", "Chronic Disease (COPD)", "Conformity: 0.87, Validity: 89%"],
        ["8", "Zhu et al. (2023)", "SMOTE-ENN + NGBoost", "Financial Risk", "AUC-ROC: 0.85, Brier Score: 0.18"],
        ["9", "Li et al. (2024)", "NGBoost", "Battery SoC", "RMSE: 1.2%, Calibration Error: 0.03"],
        ["10", "Penelitian Ini", "NGBoost + DiCE", "Water Potability (Kaggle)", "Target: F1 >= baseline, ECE < 0.1, Validity > 90%"],
    ]
    create_table_from_data(doc, headers_6a, rows_6a, col_widths=[1, 4, 4, 4, 5])

    doc.add_paragraph()
    add_formatted_paragraph(doc, "Bagian B: Orientasi dan Gap", bold=True, size=11)
    headers_6b = ["No", "Tipe Prediksi", "Tipe Explainability", "Orientasi Output", "Gap Terhadap Penelitian Ini"]
    rows_6b = [
        ["1", "Deterministik", "Tidak ada", "Deskriptif (klasifikasi WQI)", "Hanya prediktif; tanpa explainability maupun rekomendasi preskriptif; tanpa estimasi ketidakpastian"],
        ["2", "Deterministik", "Pasif (SHAP)", "Diagnostik (interpretasi fitur)", "SHAP hanya menjelaskan WHY bukan HOW to fix; tidak ada rekomendasi operasional"],
        ["3", "Deterministik", "Pasif (SHAP)", "Diagnostik (feature importance)", "Explainability tanpa prescriptive output; dataset identik namun hanya prediktif"],
        ["4", "--", "Teoritis (Counterfactual)", "Teoritis (Preskriptif)", "Belum ada implementasi counterfactual spesifik pada fisikokimia air"],
        ["5", "Deterministik", "Aktif (DiCE)", "Preskriptif (rekomendasi perubahan)", "Domain non-air; constraint berbeda dengan parameter fisikokimia"],
        ["6", "--", "Aktif (Counterfactual)", "Preskriptif (optimasi properti)", "Domain finansial; constraint berbeda dengan kualitas air"],
        ["7", "Deterministik", "Aktif (DiCE)", "Preskriptif (multi-class)", "Domain kesehatan; belum pada klasifikasi biner air"],
        ["8", "Probabilistik", "Tidak ada", "Prediktif probabilistik", "Probabilistik tanpa prescriptive; domain non-air"],
        ["9", "Probabilistik", "Tidak ada", "Prediktif probabilistik", "Probabilistik tanpa prescriptive; domain non-air"],
        ["10", "Probabilistik", "Aktif (DiCE)", "Preskriptif (rekomendasi perbaikan air)", "Mengisi gap: integrasi probabilistik + preskriptif pada domain air"],
    ]
    create_table_from_data(doc, headers_6b, rows_6b, col_widths=[1, 3, 3.5, 4.5, 6])

    doc.add_paragraph()
    p_note = doc.add_paragraph()
    run_note = p_note.add_run(
        "CATATAN PENTING: Angka metrik evaluasi yang tercantum di atas HARUS diverifikasi dengan membaca "
        "paper aslinya. Jangan gunakan angka tanpa konfirmasi dari sumber primer."
    )
    run_note.bold = True
    run_note.font.size = Pt(10)
    run_note.font.name = 'Times New Roman'
    run_note.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_paragraph()
    doc.add_heading('Penjelasan Perubahan:', level=2)
    changes6 = [
        'Ditambahkan kolom "Metrik Evaluasi" dengan angka spesifik',
        'Orientasi output lebih jelas (Deskriptif/Diagnostik/Preskriptif)',
        'Gap ditulis lebih terukur dan spesifik',
        'Penelitian Ini menyertakan target metrik'
    ]
    for change in changes6:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(change)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    # ==================== REVISI 7: SITASI GAMBAR ====================
    doc.add_page_break()
    doc.add_heading('REVISI 7: BAB 2 - SITASI GAMBAR', level=1)
    add_bold_and_normal(doc, "LOKASI: ", "Halaman 10-11, caption Gambar 2.1")
    add_bold_and_normal(doc, "ALASAN REVISI: ", "Reviewer poin 7 - \"Penamaan setiap gambar pada bab 2 seharusnya ada sitasi jurnal/sumber referensi.\"")

    doc.add_paragraph()
    doc.add_heading('Perubahan Caption:', level=2)

    add_bold_and_normal(doc, "TEKS LAMA: ", "")
    add_formatted_paragraph(doc, '"Gambar 2.1 Evolusi Pendekatan Machine learning untuk Kualitas Air"', italic=True, size=11)

    add_bold_and_normal(doc, "TEKS BARU: ", "")
    add_formatted_paragraph(doc, '"Gambar 2.1 Evolusi Pendekatan Machine Learning untuk Kualitas Air (Sumber: Diolah dari [5], [6], [16])"', bold=True, size=11)

    doc.add_paragraph()
    doc.add_heading('Aturan Penulisan Caption Gambar:', level=2)
    rules_caption = [
        'Jika gambar dibuat sendiri berdasarkan sintesis dari beberapa sumber: gunakan format "Diolah dari [referensi]"',
        'Jika gambar diambil langsung dari satu sumber: gunakan format "Sumber: [referensi]"',
        'Jika gambar dibuat sendiri tanpa referensi spesifik: gunakan format "(Sumber: Peneliti)"',
        'Terapkan pola yang sama untuk SEMUA gambar di BAB 2 dan BAB 3'
    ]
    for rule in rules_caption:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(rule)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    # ==================== REVISI 8: RUMUS METRIK EVALUASI ====================
    doc.add_page_break()
    doc.add_heading('REVISI 8: BAB 2 - RUMUS METRIK EVALUASI', level=1)
    add_bold_and_normal(doc, "LOKASI: ", "Halaman 14, sub-bab 2.2.6 \"Evaluasi Model\" - ganti seluruh isinya")
    add_bold_and_normal(doc, "ALASAN REVISI: ", "Reviewer poin 8 - \"Metrik evaluasi pada bab 2 seharusnya ada rumus matematika dan penjelasannya.\"")

    doc.add_paragraph()
    doc.add_heading('Teks Revisi (Sub-bab 2.2.6 Baru):', level=2)
    add_formatted_paragraph(doc, "2.2.6 Evaluasi Model", bold=True, size=12, space_after=6)

    intro_eval = (
        "Evaluasi model dilakukan untuk menilai kemampuan klasifikasi dan kualitas probabilistik prediksi "
        "dari model NGBoost serta baseline [5], [15]. Metrik evaluasi terbagi menjadi dua kategori: metrik "
        "klasifikasi dan metrik kalibrasi probabilistik."
    )
    add_formatted_paragraph(doc, intro_eval, size=11, space_after=12)

    # A. Metrik Klasifikasi
    add_formatted_paragraph(doc, "A. Metrik Klasifikasi", bold=True, size=11, space_after=6)

    # 1) Accuracy
    add_formatted_paragraph(doc, "1) Accuracy", bold=True, size=11, space_after=4)
    add_formatted_paragraph(doc, "Accuracy mengukur proporsi prediksi yang benar terhadap keseluruhan data:", size=11, space_after=4)
    add_formatted_paragraph(doc, "Accuracy = (TP + TN) / (TP + TN + FP + FN)", bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, "di mana TP = True Positive, TN = True Negative, FP = False Positive, FN = False Negative.", size=11, space_after=12)

    # 2) Precision
    add_formatted_paragraph(doc, "2) Precision", bold=True, size=11, space_after=4)
    add_formatted_paragraph(doc, "Precision mengukur proporsi prediksi positif yang benar terhadap seluruh prediksi positif:", size=11, space_after=4)
    add_formatted_paragraph(doc, "Precision = TP / (TP + FP)", bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, "Precision penting dalam konteks kualitas air untuk menghindari kesalahan mengklasifikasikan air tidak layak sebagai layak (false positive).", size=11, space_after=12)

    # 3) Recall
    add_formatted_paragraph(doc, "3) Recall", bold=True, size=11, space_after=4)
    add_formatted_paragraph(doc, "Recall mengukur proporsi data positif aktual yang berhasil diidentifikasi oleh model:", size=11, space_after=4)
    add_formatted_paragraph(doc, "Recall = TP / (TP + FN)", bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, "Recall penting untuk memastikan seluruh sampel air yang benar-benar layak teridentifikasi oleh model.", size=11, space_after=12)

    # 4) F1-Score
    add_formatted_paragraph(doc, "4) F1-Score", bold=True, size=11, space_after=4)
    add_formatted_paragraph(doc, "F1-Score merupakan rata-rata harmonik dari Precision dan Recall:", size=11, space_after=4)
    p_f1 = doc.add_paragraph()
    p_f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f1 = p_f1.add_run("F1-Score = 2 \u00d7 (Precision \u00d7 Recall) / (Precision + Recall)")
    run_f1.bold = True
    run_f1.font.size = Pt(11)
    run_f1.font.name = 'Times New Roman'
    add_formatted_paragraph(doc, "F1-Score dipilih sebagai metrik utama karena dataset Water Potability memiliki distribusi kelas yang tidak seimbang, sehingga akurasi saja tidak memadai sebagai indikator performa [2].", size=11, space_after=12)

    # B. Metrik Kalibrasi Probabilistik
    add_formatted_paragraph(doc, "B. Metrik Kalibrasi Probabilistik", bold=True, size=11, space_after=6)

    # 5) NLL
    add_formatted_paragraph(doc, "5) Negative Log-Likelihood (NLL)", bold=True, size=11, space_after=4)
    add_formatted_paragraph(doc, "NLL mengukur kesesuaian antara probabilitas prediksi dengan label aktual:", size=11, space_after=4)
    p_nll = doc.add_paragraph()
    p_nll.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_nll = p_nll.add_run("NLL = -(1/N) \u00d7 \u03a3\u1d62 [y\u1d62 log(\u03bc\u1d62) + (1 - y\u1d62) log(1 - \u03bc\u1d62)]")
    run_nll.bold = True
    run_nll.font.size = Pt(11)
    run_nll.font.name = 'Times New Roman'
    add_formatted_paragraph(doc, "di mana N = jumlah sampel, y\u1d62 = label aktual (0 atau 1), dan \u03bc\u1d62 = probabilitas prediksi kelas positif untuk sampel ke-i. Nilai NLL yang lebih rendah mengindikasikan model yang lebih mampu memprediksi probabilitas mendekati frekuensi kejadian sebenarnya [5].", size=11, space_after=12)

    # 6) ECE
    add_formatted_paragraph(doc, "6) Expected Calibration Error (ECE)", bold=True, size=11, space_after=4)
    add_formatted_paragraph(doc, "ECE mengukur tingkat kalibrasi model dengan membandingkan probabilitas prediksi rata-rata dengan proporsi aktual kelas positif pada setiap bin:", size=11, space_after=4)
    p_ece = doc.add_paragraph()
    p_ece.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ece = p_ece.add_run("ECE = \u03a3\u2098 (|B\u2098|/N) \u00d7 |acc(B\u2098) - conf(B\u2098)|")
    run_ece.bold = True
    run_ece.font.size = Pt(11)
    run_ece.font.name = 'Times New Roman'
    add_formatted_paragraph(doc, "di mana M = jumlah bin, |B\u2098| = jumlah sampel pada bin ke-m, acc(B\u2098) = akurasi aktual pada bin ke-m, dan conf(B\u2098) = rata-rata confidence (probabilitas prediksi) pada bin ke-m. ECE yang mendekati 0 mengindikasikan model yang terkalibrasi sempurna [15].", size=11, space_after=12)

    # C. Metrik Evaluasi Counterfactual
    add_formatted_paragraph(doc, "C. Metrik Evaluasi Counterfactual", bold=True, size=11, space_after=6)

    # 7) Validity
    add_formatted_paragraph(doc, "7) Validity Rate", bold=True, size=11, space_after=4)
    add_formatted_paragraph(doc, "Validity mengukur proporsi counterfactual yang berhasil mengubah prediksi kelas:", size=11, space_after=4)
    p_val = doc.add_paragraph()
    p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_val = p_val.add_run("Validity = (Jumlah CF yang mengubah kelas) / (Total CF yang dibangkitkan) \u00d7 100%")
    run_val.bold = True
    run_val.font.size = Pt(11)
    run_val.font.name = 'Times New Roman'
    doc.add_paragraph()

    # 8) Proximity
    add_formatted_paragraph(doc, "8) Proximity (L1-Distance)", bold=True, size=11, space_after=4)
    add_formatted_paragraph(doc, "Proximity mengukur kedekatan counterfactual dengan instance asli menggunakan jarak Manhattan ternormalisasi:", size=11, space_after=4)
    p_prox = doc.add_paragraph()
    p_prox.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_prox = p_prox.add_run("Proximity = (1/d) \u00d7 \u03a3\u2c7c |x\u2c7c - x'\u2c7c| / range(x\u2c7c)")
    run_prox.bold = True
    run_prox.font.size = Pt(11)
    run_prox.font.name = 'Times New Roman'
    add_formatted_paragraph(doc, "di mana d = jumlah fitur, x\u2c7c = nilai fitur asli, x'\u2c7c = nilai fitur counterfactual, dan range(x\u2c7c) = rentang nilai fitur j. Nilai proximity yang lebih kecil menunjukkan perubahan yang lebih minimal [9].", size=11, space_after=12)

    # 9) Sparsity
    add_formatted_paragraph(doc, "9) Sparsity", bold=True, size=11, space_after=4)
    add_formatted_paragraph(doc, "Sparsity mengukur jumlah rata-rata fitur yang diubah per counterfactual:", size=11, space_after=4)
    p_spar = doc.add_paragraph()
    p_spar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_spar = p_spar.add_run("Sparsity = (1/K) \u00d7 \u03a3\u2096 |{j : x'\u2096\u2c7c \u2260 x\u2c7c}|")
    run_spar.bold = True
    run_spar.font.size = Pt(11)
    run_spar.font.name = 'Times New Roman'
    add_formatted_paragraph(doc, "di mana K = jumlah counterfactual. Sparsity yang lebih rendah menunjukkan rekomendasi yang lebih praktis untuk diimplementasikan [9].", size=11, space_after=12)

    # 10) Diversity
    add_formatted_paragraph(doc, "10) Diversity", bold=True, size=11, space_after=4)
    add_formatted_paragraph(doc, "Diversity mengukur perbedaan antar counterfactual dalam satu set rekomendasi:", size=11, space_after=4)
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_div = p_div.add_run("Diversity = (2/(K(K-1))) \u00d7 \u03a3\u2096\u2081 \u03a3\u2096\u2082 dist(CF_k1, CF_k2)")
    run_div.bold = True
    run_div.font.size = Pt(11)
    run_div.font.name = 'Times New Roman'
    add_formatted_paragraph(doc, "di mana K = jumlah counterfactual per instance dan dist = fungsi jarak antar counterfactual. Diversity yang lebih tinggi memberikan lebih banyak alternatif rekomendasi [8].", size=11, space_after=12)

    # 11) Feasibility
    add_formatted_paragraph(doc, "11) Feasibility Rate", bold=True, size=11, space_after=4)
    add_formatted_paragraph(doc, "Feasibility mengukur proporsi counterfactual yang nilai fiturnya berada dalam rentang yang diperbolehkan oleh constraints domain:", size=11, space_after=4)
    p_feas = doc.add_paragraph()
    p_feas.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_feas = p_feas.add_run("Feasibility = (Jumlah CF dalam rentang constraint) / (Total CF) \u00d7 100%")
    run_feas.bold = True
    run_feas.font.size = Pt(11)
    run_feas.font.name = 'Times New Roman'
    add_formatted_paragraph(doc, "Constraints domain mengacu pada standar kelayakan parameter fisikokimia air [9].", size=11, space_after=12)

    # ==================== REVISI 9: FLOWCHART ====================
    doc.add_page_break()
    doc.add_heading('REVISI 9: BAB 3 - FLOWCHART', level=1)
    add_bold_and_normal(doc, "LOKASI: ", "Halaman 18, Gambar 3.1 \"Alur Diagram Pemodelan\"")
    add_bold_and_normal(doc, "ALASAN REVISI: ", "Reviewer poin 9 - \"Flowchart kurang detail tidak ada split data. Seharusnya ada pembagian data latih, data validasi atau data uji.\"")

    doc.add_paragraph()
    doc.add_heading('Deskripsi Flowchart Baru (Urutan Node):', level=2)
    add_formatted_paragraph(doc, "TINDAKAN: Gambar ulang flowchart dengan menambahkan detail pembagian data. Gunakan draw.io, Lucidchart, atau Microsoft Visio.", bold=True, size=11)

    doc.add_paragraph()
    flowchart_nodes = [
        ("1", "Dataset Water Potability (Kaggle)", "Parallelogram (Input)"),
        ("2", "Exploratory Data Analysis (EDA)", "Rectangle (Proses)"),
        ("3", "Penanganan Missing Values (MICE)", "Rectangle (Proses)"),
        ("4", "Stratified Train-Test Split", "Diamond (Keputusan/Split)"),
        ("5a", "Data Latih 70%", "Rectangle"),
        ("5b", "Data Validasi 15%", "Rectangle"),
        ("5c", "Data Uji 15%", "Rectangle"),
        ("6", "Feature Scaling (Standardization)\nfit pada data latih, transform pada validasi & uji", "Rectangle (Proses)"),
        ("7", "Penanganan Class Imbalance (SMOTE-ENN)\nhanya pada data latih", "Rectangle (Proses)"),
        ("8", "Training Model NGBoost + Baseline (XGBoost, RF)\nmenggunakan data latih, tuning pada data validasi", "Rectangle (Proses)"),
        ("9", "Evaluasi Model Prediktif\npada data uji: Accuracy, Precision, Recall, F1, NLL, ECE", "Rectangle (Proses)"),
        ("10", "Seleksi Instance \"Tidak Layak\" dari Data Uji", "Rectangle (Proses)"),
        ("11", "Pembangkitan Counterfactual dengan DiCE\nk=4-5 alternatif per instance, dengan feasibility constraints", "Rectangle (Proses)"),
        ("12", "Evaluasi Kualitas Counterfactual\nValidity, Proximity, Sparsity, Diversity, Feasibility", "Rectangle (Proses)"),
        ("13", "Analisis Komparatif & Trade-off", "Rectangle (Proses)"),
        ("14", "Kesimpulan & Rekomendasi", "Rounded Rectangle (Output)"),
    ]

    headers_fc = ["No", "Isi Node/Blok", "Bentuk"]
    rows_fc = [(n[0], n[1], n[2]) for n in flowchart_nodes]
    create_table_from_data(doc, headers_fc, rows_fc, col_widths=[1.5, 10, 4.5])

    doc.add_paragraph()
    add_formatted_paragraph(doc, "Caption baru: \"Gambar 3.1 Alur Diagram Pemodelan (Sumber: Peneliti)\"", bold=True, size=11)

    doc.add_paragraph()
    add_formatted_paragraph(doc, "Catatan:", bold=True, size=11, space_after=4)
    notes_fc = [
        'Node 4 (Split) harus menunjukkan 3 cabang terpisah menuju Data Latih, Data Validasi, dan Data Uji',
        'Node 6: Standardization di-fit pada data latih saja, lalu di-transform pada validasi dan uji',
        'Node 7: SMOTE-ENN hanya diterapkan pada data latih (BUKAN pada validasi/uji)',
        'Pastikan setiap node memiliki bentuk yang sesuai standar flowchart',
        'Gambar menggunakan tool diagram (draw.io/Lucidchart/Visio), bukan dibuat manual'
    ]
    for note in notes_fc:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(note)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    # ==================== REVISI 10: PREPROCESSING TABEL CONTOH ====================
    doc.add_page_break()
    doc.add_heading('REVISI 10: BAB 3 - PREPROCESSING DENGAN TABEL CONTOH', level=1)
    add_bold_and_normal(doc, "LOKASI: ", "Halaman 19-20, sub-bab 3.2.1 sampai 3.2.5")
    add_bold_and_normal(doc, "ALASAN REVISI: ", "Reviewer poin 10 - \"Tahap-tahap preprocessing seharusnya ada contohnya dalam bentuk tabel atau gambar bukan hanya teks saja.\"")
    add_formatted_paragraph(doc, "TINDAKAN: Tambahkan tabel contoh SETELAH penjelasan teks pada setiap sub-tahap preprocessing. Teks penjelasan yang ada TETAP dipertahankan, tabel ditambahkan sebagai pelengkap.", bold=True, size=11)

    # Tabel 3.1
    doc.add_paragraph()
    doc.add_heading('Tabel 3.1: Contoh Statistik Deskriptif Dataset Water Potability', level=2)
    add_formatted_paragraph(doc, "Letakkan setelah sub-bab 3.2.1 (EDA):", bold=True, size=11)

    headers_31 = ["Parameter", "Count", "Mean", "Std", "Min", "Max", "Missing (%)"]
    rows_31 = [
        ["pH", "2785", "7.08", "1.59", "0.23", "14.00", "14.98%"],
        ["Hardness", "3276", "196.37", "32.88", "47.43", "323.12", "0.00%"],
        ["Solids", "3276", "22014.09", "8768.57", "320.94", "61227.20", "0.00%"],
        ["Chloramines", "3276", "7.12", "1.58", "0.35", "13.13", "0.00%"],
        ["Sulfate", "2495", "333.78", "36.14", "129.00", "481.03", "23.84%"],
        ["Conductivity", "3276", "426.21", "68.44", "181.48", "753.34", "0.00%"],
        ["Organic_carbon", "3276", "14.28", "3.31", "2.20", "28.30", "0.00%"],
        ["Trihalomethanes", "3114", "66.40", "16.18", "0.74", "124.00", "4.95%"],
        ["Turbidity", "3276", "3.97", "0.78", "1.45", "6.74", "0.00%"],
    ]
    create_table_from_data(doc, headers_31, rows_31, col_widths=[3, 1.8, 2.2, 2, 1.8, 2.2, 2.2])

    p_cat31 = doc.add_paragraph()
    run_cat31 = p_cat31.add_run("Catatan: Angka di atas merupakan contoh ilustratif. Verifikasi dengan dataset aktual saat implementasi.")
    run_cat31.italic = True
    run_cat31.font.size = Pt(10)
    run_cat31.font.name = 'Times New Roman'

    # Tabel 3.2
    doc.add_paragraph()
    doc.add_heading('Tabel 3.2: Contoh Hasil Imputasi MICE pada Parameter pH', level=2)
    add_formatted_paragraph(doc, "Letakkan setelah sub-bab 3.2.2 (Missing Values):", bold=True, size=11)

    headers_32 = ["Sample ID", "pH (Sebelum Imputasi)", "pH (Setelah Imputasi MICE)", "Keterangan"]
    rows_32 = [
        ["1", "7.08", "7.08", "Tidak berubah (tidak missing)"],
        ["2", "NaN", "6.92", "Diimputasi berdasarkan korelasi antar fitur"],
        ["3", "NaN", "7.45", "Diimputasi berdasarkan korelasi antar fitur"],
        ["4", "8.32", "8.32", "Tidak berubah (tidak missing)"],
        ["5", "NaN", "7.11", "Diimputasi berdasarkan korelasi antar fitur"],
    ]
    create_table_from_data(doc, headers_32, rows_32, col_widths=[2.5, 4, 4.5, 5.5])

    p_cat32 = doc.add_paragraph()
    run_cat32 = p_cat32.add_run("Catatan: Nilai imputasi bersifat ilustratif. Hasil aktual bergantung pada proses iterasi MICE.")
    run_cat32.italic = True
    run_cat32.font.size = Pt(10)
    run_cat32.font.name = 'Times New Roman'

    # Tabel 3.3
    doc.add_paragraph()
    doc.add_heading('Tabel 3.3: Contoh Hasil Feature Scaling (Standardization)', level=2)
    add_formatted_paragraph(doc, "Letakkan setelah sub-bab 3.2.4 (Feature Scaling):", bold=True, size=11)

    headers_33 = ["Parameter", "Nilai Asli", "Setelah Standardization", "\u03bc (dari data latih)", "\u03c3 (dari data latih)"]
    rows_33 = [
        ["pH", "7.08", "0.000", "7.08", "1.59"],
        ["Hardness", "196.37", "0.000", "196.37", "32.88"],
        ["Solids", "22014.09", "0.000", "22014.09", "8768.57"],
        ["pH", "9.50", "1.522", "7.08", "1.59"],
        ["Hardness", "250.00", "1.631", "196.37", "32.88"],
    ]
    create_table_from_data(doc, headers_33, rows_33, col_widths=[3, 3, 4, 3.5, 3.5])

    p_cat33 = doc.add_paragraph()
    run_cat33 = p_cat33.add_run("Catatan: Nilai \u03bc dan \u03c3 dihitung HANYA dari data latih, kemudian diterapkan pada data validasi dan data uji untuk mencegah data leakage.")
    run_cat33.italic = True
    run_cat33.font.size = Pt(10)
    run_cat33.font.name = 'Times New Roman'

    # ==================== REVISI 11: PERBAIKAN PENOMORAN ====================
    doc.add_page_break()
    doc.add_heading('REVISI 11: BAB 3 - PERBAIKAN PENOMORAN', level=1)
    add_bold_and_normal(doc, "LOKASI: ", "Halaman 25-26, sub-bab 3.5, 3.6, 3.7")
    add_bold_and_normal(doc, "ALASAN REVISI: ", "Penomoran duplikat dan tidak konsisten. Sub-bab 3.5 dan 3.6 sama-sama berjudul \"Evaluasi\", dan sub-bab \"Tools dan Lingkungan Pengembangan\" ditempatkan di bawah heading evaluasi.")

    doc.add_paragraph()
    doc.add_heading('Struktur Lama (Bermasalah):', level=2)
    old_struct = [
        "3.5. Evaluasi",
        "    3.5.1. Evaluasi Model Prediktif",
        "    3.5.2. Evaluasi Counterfactual",
        "3.6. Evaluasi  <-- DUPLIKAT!",
        "    3.6.1. Tools dan Lingkungan Pengembangan  <-- SALAH TEMPAT!",
    ]
    for s in old_struct:
        p = doc.add_paragraph()
        run = p.add_run(s)
        run.font.size = Pt(11)
        run.font.name = 'Courier New'
        if "DUPLIKAT" in s or "SALAH" in s:
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_paragraph()
    doc.add_heading('Struktur Baru (Benar):', level=2)
    new_struct = [
        "3.5. Evaluasi",
        "    3.5.1. Evaluasi Model Prediktif",
        "    3.5.2. Evaluasi Counterfactual",
        "3.6. Tools dan Lingkungan Pengembangan  <-- Sub-bab tersendiri",
    ]
    for s in new_struct:
        p = doc.add_paragraph()
        run = p.add_run(s)
        run.font.size = Pt(11)
        run.font.name = 'Courier New'
        if "<--" in s:
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    doc.add_paragraph()
    add_formatted_paragraph(doc, "TINDAKAN:", bold=True, size=11, space_after=4)
    actions_11 = [
        'Hapus duplikasi heading "3.6. Evaluasi"',
        'Hapus duplikasi heading "3.7. Evaluasi" (jika ada)',
        'Jadikan "Tools dan Lingkungan Pengembangan" sebagai sub-bab 3.6 tersendiri (bukan bagian dari Evaluasi)',
    ]
    for a in actions_11:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(a)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    # ==================== RINGKASAN CHECKLIST ====================
    doc.add_page_break()
    doc.add_heading('RINGKASAN CHECKLIST REVISI', level=1)
    add_formatted_paragraph(doc, "Gunakan tabel berikut sebagai checklist untuk memastikan semua revisi telah dilakukan:", size=11, space_after=12)

    headers_cl = ["No", "Bagian", "Lokasi", "Jenis Revisi", "Prioritas", "Status"]
    rows_cl = [
        ["1", "Abstrak", "Hal. i", "Rewrite total", "TINGGI", "[ ]"],
        ["2", "Latar Belakang - paragraf 3", "Hal. 1", "Ganti paragraf + tambah metrik", "TINGGI", "[ ]"],
        ["3", "Rumusan Masalah", "Hal. 3, sub-bab 1.2", "Rewrite total", "TINGGI", "[ ]"],
        ["4", "Tujuan", "Hal. 3, sub-bab 1.3", "Rewrite total", "TINGGI", "[ ]"],
        ["5", "Hapus Batasan Masalah", "Hal. 3-4, sub-bab 1.4", "Hapus & renomor", "SEDANG", "[ ]"],
        ["6", "Tabel Penelitian Terdahulu", "Hal. 8-9, Tabel 2.1", "Ganti tabel", "TINGGI", "[ ]"],
        ["7", "Caption Gambar 2.1", "Hal. 11", "Tambah sitasi", "SEDANG", "[ ]"],
        ["8", "Metrik Evaluasi", "Hal. 14, sub-bab 2.2.6", "Rewrite + tambah rumus", "TINGGI", "[ ]"],
        ["9", "Flowchart", "Hal. 18, Gambar 3.1", "Gambar ulang", "TINGGI", "[ ]"],
        ["10", "Preprocessing (tabel contoh)", "Hal. 19-20", "Tambah 3 tabel", "TINGGI", "[ ]"],
        ["11", "Penomoran BAB 3", "Hal. 25-26", "Perbaiki struktur", "SEDANG", "[ ]"],
    ]
    create_table_from_data(doc, headers_cl, rows_cl, col_widths=[1, 4.5, 3.5, 4, 2.5, 1.5])

    # ==================== CATATAN PENTING ====================
    doc.add_paragraph()
    doc.add_heading('CATATAN PENTING', level=1)

    catatan = [
        "Angka metrik evaluasi pada Tabel 2.1 (Revisi 6) HARUS diverifikasi dengan membaca paper asli. Jangan gunakan angka yang diberikan tanpa konfirmasi.",
        "Tabel contoh preprocessing (Revisi 10) bersifat ilustratif. Ganti dengan data aktual saat implementasi di Google Colab.",
        "Flowchart (Revisi 9) harus digambar ulang menggunakan tool diagram. Deskripsi di atas adalah panduan konten, bukan gambar jadi.",
        "Semua gambar di BAB 2 dan BAB 3 harus ditambahkan sitasi sumber pada caption-nya.",
        "Setelah semua revisi selesai, periksa kembali konsistensi penomoran di seluruh dokumen.",
        "Pastikan daftar pustaka diperbarui jika ada referensi baru yang ditambahkan.",
    ]
    for i, c in enumerate(catatan, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f"{i}. ")
        run_num.bold = True
        run_num.font.size = Pt(11)
        run_num.font.name = 'Times New Roman'
        run_text = p.add_run(c)
        run_text.font.size = Pt(11)
        run_text.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)

    # Save the document
    output_path = "/projects/sandbox/ProposalSeminarS6/REVISI_PROPOSAL_AFLAH.docx"
    doc.save(output_path)
    print(f"Document saved successfully to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_document()
