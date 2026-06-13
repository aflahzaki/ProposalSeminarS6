"""Build REVISI_PROPOSAL_AFLAH.docx v6 - Final revision guide (28 revisions)."""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # headers
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    # data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
    # borders
    tbl = t._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml('<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    borders = parse_xml(
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    doc.add_paragraph()
    return t

# === HALAMAN JUDUL ===
p = doc.add_heading('PANDUAN REVISI PROPOSAL TUGAS AKHIR (v6 - Final)', level=1)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('"Analisis Preskriptif Kualitas Air Menggunakan Natural Gradient Boosting dan Counterfactual Explanations"')
r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Nama: Aflah Zaki Siregar | NIM: 103062300095')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Dokumen panduan revisi final (28 revisi). Versi 6 menambahkan perbaikan BAB 2: justifikasi metode NGBoost dan DiCE, tabel perbandingan XAI, batasan konsep probabilistik, what-if analysis, dan novelitas statement.')

doc.add_page_break()

# === REVISI 1 ===
doc.add_heading('REVISI 1: ABSTRAK', level=2)
doc.add_paragraph('Lokasi: Halaman i, ganti seluruh paragraf dan kata kunci')
doc.add_paragraph('Alasan: Reviewer poin 1,2 + Template TA (hipotesa awal)')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Teks Revisi:').bold = True
doc.add_paragraph(
    'Kualitas air minum yang memenuhi standar kelayakan merupakan aspek fundamental dalam menjaga kesehatan masyarakat. '
    'Penelitian terdahulu menunjukkan bahwa berbagai algoritma machine learning seperti XGBoost dan Random Forest mampu mengklasifikasikan kelayakan air dengan akurasi tinggi, '
    'namun pendekatan-pendekatan tersebut hanya menghasilkan label prediksi tanpa menyertakan estimasi ketidakpastian maupun rekomendasi tindakan perbaikan yang dapat ditindaklanjuti secara operasional. '
    'Kesenjangan antara kemampuan prediktif dan kebutuhan preskriptif di lapangan menjadi urgensi utama yang belum terjawab oleh literatur yang ada. '
    'Penelitian ini bertujuan untuk menganalisis kerangka kerja analisis preskriptif yang mengintegrasikan Natural Gradient Boosting (NGBoost) dengan Diverse Counterfactual Explanations (DiCE) pada klasifikasi kelayakan air minum. '
    'NGBoost digunakan untuk memodelkan distribusi probabilitas kelayakan air melalui distribusi Bernoulli dan Natural Gradient, '
    'sedangkan DiCE diterapkan pada model terlatih untuk membangkitkan rekomendasi actionable recourse bagi sampel tidak layak dengan mempertimbangkan properti validity, proximity, sparsity, diversity, dan feasibility. '
    'Analisis komparatif dilakukan terhadap baseline deterministik (XGBoost dan Random Forest) menggunakan dataset Water Potability publik dengan evaluasi metrik klasifikasi (Accuracy, Precision, Recall, F1-Score) dan kalibrasi probabilitas (NLL, ECE). '
    'Dihipotesiskan bahwa NGBoost mampu menghasilkan performa klasifikasi setara atau lebih baik dibandingkan baseline deterministik dengan kalibrasi probabilistik yang superior, '
    'serta DiCE mampu menghasilkan rekomendasi counterfactual dengan validity rate di atas 90% dan feasibility rate di atas 85% sesuai constraint domain kualitas air.'
)

p = doc.add_paragraph()
p.add_run('Kata Kunci: ').bold = True
p.add_run('analisis preskriptif, Natural Gradient Boosting, counterfactual explanations, kualitas air minum, prediksi probabilistik, actionable recourse')

doc.add_page_break()

# === REVISI 2 ===
doc.add_heading('REVISI 2: LATAR BELAKANG', level=2)
doc.add_paragraph('Lokasi: Paragraf "Untuk mengatasi kompleksitas data tersebut..."')
doc.add_paragraph('Alasan: Reviewer poin 1,3')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Teks Revisi:').bold = True
doc.add_paragraph(
    'Untuk mengatasi kompleksitas data tersebut, berbagai algoritma machine learning berbasis ensemble telah banyak diterapkan dalam pemantauan kualitas air. '
    'Park et al. (2022) mengimplementasikan ensemble learning dengan SHAP dan mencapai akurasi ~80% [6]. '
    'Patel et al. (2022) menerapkan RF, XGBoost, Gradient Boosting pada dataset Water Potability dari Kaggle yang identik dengan penelitian ini, memperoleh akurasi tertinggi 71% dengan Random Forest [NEW_REF_PATEL]. '
    'Zhu et al. (2023) mengintegrasikan SMOTE-ENN dengan NGBoost dan menunjukkan AUC-ROC 0.85 [4]. '
    'Meskipun algoritma tersebut mampu menghasilkan akurasi kompetitif, keseluruhan pendekatan masih berhenti pada tahap prediktif/diagnostik. '
    'Pendekatan explainability post-hoc seperti SHAP dan LIME hanya menjelaskan alasan di balik prediksi, namun tidak memberikan instruksi bagaimana parameter harus diubah untuk memperbaiki kualitas air. '
    'Kesenjangan fundamental ini bukan masalah transparansi model, melainkan ketiadaan mekanisme preskriptif yang mampu mentransformasi output prediksi menjadi rekomendasi tindakan operasional.'
)

doc.add_page_break()

# === REVISI 3 ===
doc.add_heading('REVISI 3: RUMUSAN MASALAH', level=2)
doc.add_paragraph('Lokasi: Seluruh isi sub-bab 1.2')
doc.add_paragraph('Alasan: Reviewer poin 4')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Teks Revisi:').bold = True
doc.add_paragraph('Rumusan masalah:')
doc.add_paragraph('1. Bagaimana performa NGBoost dalam memodelkan kelayakan air secara probabilistik dibandingkan baseline deterministik (XGBoost, RF) berdasarkan metrik klasifikasi dan kalibrasi?')
doc.add_paragraph('2. Bagaimana implementasi DiCE pada model NGBoost untuk menghasilkan rekomendasi preskriptif yang mengubah status air dari tidak layak menjadi layak?')
doc.add_paragraph('3. Sejauh mana kualitas rekomendasi counterfactual memenuhi properti validity, proximity, sparsity, diversity, feasibility, serta bagaimana trade-off antar properti tersebut?')

doc.add_page_break()

# === REVISI 4 ===
doc.add_heading('REVISI 4: TUJUAN', level=2)
doc.add_paragraph('Lokasi: Seluruh isi sub-bab 1.3')
doc.add_paragraph('Alasan: Reviewer poin 4')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Teks Revisi:').bold = True
doc.add_paragraph('Tujuan khusus:')
doc.add_paragraph('1. Menganalisis performa NGBoost secara komparatif terhadap baseline (XGBoost, RF) dari metrik klasifikasi (Accuracy, Precision, Recall, F1) dan kalibrasi (NLL, ECE).')
doc.add_paragraph('2. Mengimplementasikan DiCE pada model NGBoost untuk menghasilkan rekomendasi perubahan parameter fisikokimia yang mengubah status air.')
doc.add_paragraph('3. Menganalisis kualitas counterfactual berdasarkan validity, proximity, sparsity, diversity, feasibility, termasuk trade-off antar properti dalam constraint domain kualitas air.')

doc.add_page_break()

# === REVISI 5 ===
doc.add_heading('REVISI 5: HAPUS BATASAN MASALAH', level=2)
doc.add_paragraph('Lokasi: Sub-bab 1.4. Tindakan: Hapus. Renomor 1.5->1.4, 1.6->1.5.')

doc.add_page_break()

# === REVISI 6 ===
doc.add_heading('REVISI 6: TABEL PENELITIAN TERDAHULU', level=2)
doc.add_paragraph('Lokasi: Tabel 2.1. Alasan: Reviewer poin 6.')
doc.add_paragraph()

headers6 = ['No','Author','Metode','Dataset','Metrik','Prediksi','XAI','Output','Gap']
rows6 = [
    ['1','Aslam(2022)','Hybrid NN+XGB','Indus River','Acc:95.2%','Deterministik','Tidak ada','Deskriptif','Tanpa XAI'],
    ['2','Park(2022)','Ensemble+SHAP','Water Quality','Acc:~80%','Deterministik','Pasif(SHAP)','Diagnostik','WHY not HOW'],
    ['3','Patel(2022)','RF,XGB,GB','Water Potability Kaggle','Acc:71%','Deterministik','Tidak ada','Deskriptif','Identik;prediktif saja'],
    ['4','Aderemi(2025)','Sys.Review','Water Quality','-','-','Teoritis(CF)','Teoritis','Belum implementasi'],
    ['5','Nnadi(2026)','Multi-XAI+DiCE','Depression','Validity:94%','Deterministik','Aktif(DiCE)','Preskriptif','Domain non-air'],
    ['6','Dastile&Celik(2024)','CF Optim.','Credit','Validity:91%','-','Aktif(CF)','Preskriptif','Domain finansial'],
    ['7','Lenatti(2025)','DiCE vs MUCH','COPD','Conformity:0.87','Deterministik','Aktif(DiCE)','Preskriptif','Domain kesehatan'],
    ['8','Zhu(2023)','SMOTE+NGBoost','Financial','AUC:0.85','Probabilistik','Tidak ada','Pred.prob.','Tanpa preskriptif'],
    ['9','Duan(2020)','NGBoost','Benchmark','NLL,CRPS','Probabilistik','Tidak ada','Pred.prob.','Foundational'],
    ['10','Penelitian Ini','NGBoost+DiCE','Water Potability','Target:F1>=base ECE<0.1 Val>90%','Probabilistik','Aktif(DiCE)','Preskriptif','Integrasi di air'],
]
add_table(doc, headers6, rows6)

doc.add_page_break()

# === REVISI 7 ===
doc.add_heading('REVISI 7: SITASI GAMBAR', level=2)
doc.add_paragraph('Ubah caption Gambar 2.1 menjadi:')
doc.add_paragraph('"Gambar 2.1 Evolusi Pendekatan ML untuk Kualitas Air (Sumber: Diolah dari [6], [16])"')

doc.add_page_break()

# === REVISI 8 ===
doc.add_heading('REVISI 8: RUMUS METRIK EVALUASI', level=2)
doc.add_paragraph('Lokasi: Seluruh sub-bab 2.2.6. Ganti dengan rumus:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('A. Klasifikasi:').bold = True
doc.add_paragraph('Accuracy = (TP+TN) / (TP+TN+FP+FN)')
doc.add_paragraph('Precision = TP / (TP+FP)')
doc.add_paragraph('Recall = TP / (TP+FN)')
doc.add_paragraph('F1 = 2*P*R / (P+R)')

p = doc.add_paragraph()
p.add_run('B. Kalibrasi:').bold = True
doc.add_paragraph('NLL = -(1/N) * SUM[y*log(mu) + (1-y)*log(1-mu)]')
doc.add_paragraph('ECE = SUM(|Bm|/N) * |acc - conf|')

p = doc.add_paragraph()
p.add_run('C. Counterfactual:').bold = True
doc.add_paragraph('Validity = jumlah CF yang mengubah prediksi / total CF')
doc.add_paragraph('Proximity = (1/k) * SUM distance(x, cf_i)')
doc.add_paragraph('Sparsity = (1/k) * SUM (fitur berubah / total fitur)')
doc.add_paragraph('Diversity = (1/k(k-1)) * SUM distance(cf_i, cf_j)')
doc.add_paragraph('Feasibility = jumlah CF dalam range constraint / total CF')

doc.add_page_break()

# === REVISI 9 ===
doc.add_heading('REVISI 9: FLOWCHART', level=2)
doc.add_paragraph('Lokasi: Gambar 3.1. Alasan: Reviewer poin 9 + data leakage fix + baseline.')
doc.add_paragraph('Spesifikasi 27 node:')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('PREPROCESSING:').bold = True

headers9 = ['No','Node','Shape','Keterangan','Justifikasi']
rows9_pre = [
    ['1','START','Rounded Rect','-','-'],
    ['2','Water Quality Dataset','Parallelogram','Input','Patel(2022)'],
    ['3','EDA','Rectangle','Statistik, distribusi','Standard'],
    ['4','Imputation(MICE)','Rectangle','Missing values','Barrabes(2025)'],
    ['5','Stratified Split 70/15/15','Diamond','Train/Val/Test','Reviewer poin 9'],
    ['6','Scaling(StandardScaler)','Rectangle','fit TRAIN, transform val/test','Prevent leakage'],
    ['7','Imbalance Check','Diamond','-','-'],
    ['8','SMOTE-ENN(train only)','Rectangle','Jika imbalance','Zhu(2023)'],
]
add_table(doc, headers9, rows9_pre)

p = doc.add_paragraph()
p.add_run('FASE A - MODEL TRAINING:').bold = True

rows9_a = [
    ['9','Training NGBoost','Rectangle','Bernoulli, NatGrad','Duan(2020)'],
    ['10','Training Baseline(XGB,RF)','Rectangle','Parallel','Patel(2022)'],
    ['11','Hyperparam Tuning','Rectangle','Via validation','-'],
    ['12','Convergence?','Diamond','Early stopping','-'],
    ['13','Calibration(ECE)?','Diamond','ECE<threshold?','Li(2024)'],
    ['14','Recalibration','Rectangle','Platt/Isotonic','-'],
    ['15','Evaluasi Test Set','Rectangle','Acc,Prec,Rec,F1,NLL,ECE','-'],
    ['16','Analisis Komparatif','Rectangle','NGBoost vs XGB vs RF','Judul "Analisis"'],
    ['17','Seleksi Tidak Layak','Rectangle','Dari test set','Nnadi(2026)'],
]
add_table(doc, headers9, rows9_a)

p = doc.add_paragraph()
p.add_run('FASE B - COUNTERFACTUAL:').bold = True

rows9_b = [
    ['18','Input+Constraints','Rectangle','Class 0+Permenkes/WHO','-'],
    ['19','DiCE Generation(k=4-5)','Rectangle','Generate CF','Nnadi(2026)'],
    ['20','Validity Check','Diamond','>90%?','-'],
    ['21','Adjust Weights','Rectangle','Jika rendah','-'],
    ['22','Feasibility Check','Diamond','Dalam range?','-'],
    ['23','Relax Constraints','Rectangle','Jika rendah','-'],
    ['24','Evaluasi CF','Rectangle','5 properti','Dastile(2024)'],
    ['25','Analisis Trade-off','Rectangle','-','-'],
    ['26','Output Prescriptive','Rectangle','Hasil akhir','-'],
    ['27','END','Rounded Rect','-','-'],
]
add_table(doc, headers9, rows9_b)

p = doc.add_paragraph()
p.add_run('CATATAN URUTAN KRITIS: ').bold = True
p.add_run('Imputation -> Split -> Scaling -> SMOTE (jangan terbalik!)')

doc.add_page_break()

# === REVISI 10 ===
doc.add_heading('REVISI 10: TABEL PREPROCESSING', level=2)
doc.add_paragraph('3 tabel contoh (statistik deskriptif 9 param, MICE 5 row, standardization 5 row) - bersifat ilustratif.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Tabel A: Statistik Deskriptif (Ilustratif)').bold = True
headers10a = ['Parameter','Mean','Std','Min','Max','Missing(%)']
rows10a = [
    ['pH','7.08','1.59','0.23','14.00','14.98%'],
    ['Hardness','196.37','32.88','47.43','323.12','0%'],
    ['Solids (TDS)','22014.09','8768.57','320.94','61227.20','0%'],
    ['Chloramines','7.12','1.58','0.35','13.13','0%'],
    ['Sulfate','333.78','41.42','129.00','481.03','23.84%'],
    ['Conductivity','426.21','68.39','181.48','753.34','0%'],
    ['Organic Carbon','14.28','3.31','2.20','28.30','0%'],
    ['Trihalomethanes','66.40','16.18','0.74','124.00','4.95%'],
    ['Turbidity','3.97','0.78','1.45','6.74','0%'],
]
add_table(doc, headers10a, rows10a)

p = doc.add_paragraph()
p.add_run('Tabel B: Hasil Imputasi MICE (5 sampel ilustratif)').bold = True
headers10b = ['Index','pH(before)','pH(after)','Sulfate(before)','Sulfate(after)']
rows10b = [
    ['0','NaN','7.12','NaN','334.56'],
    ['5','NaN','6.89','320.50','320.50'],
    ['12','8.10','8.10','NaN','341.22'],
    ['23','NaN','7.45','NaN','328.90'],
    ['45','6.50','6.50','280.10','280.10'],
]
add_table(doc, headers10b, rows10b)

p = doc.add_paragraph()
p.add_run('Tabel C: Standardization (5 sampel ilustratif)').bold = True
headers10c = ['Parameter','Original','Scaled(z-score)']
rows10c = [
    ['pH=7.08','7.08','0.00'],
    ['pH=8.50','8.50','0.89'],
    ['Hardness=196','196.00','-0.01'],
    ['TDS=22014','22014.00','0.00'],
    ['Turbidity=5.0','5.00','1.32'],
]
add_table(doc, headers10c, rows10c)

doc.add_page_break()

# === REVISI 11 ===
doc.add_heading('REVISI 11: PENOMORAN BAB 3', level=2)
doc.add_paragraph('Struktur baru:')
doc.add_paragraph('3.5 Evaluasi')
doc.add_paragraph('  3.5.1 Model')
doc.add_paragraph('  3.5.2 CF')
doc.add_paragraph('3.6 Tools')

doc.add_page_break()

# === REVISI 12 ===
doc.add_heading('REVISI 12: NARASI DATASET', level=2)
doc.add_paragraph('Lokasi: Sub-bab 3.2')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Teks Revisi:').bold = True
doc.add_paragraph(
    '...Water Potability Dataset dari Kaggle berlisensi CC0: Public Domain. '
    'Dataset telah digunakan dalam 99+ publikasi Scopus termasuk Patel et al. (2022) dan Park et al. (2022). '
    'Dataset diposisikan sebagai benchmark validasi metodologi, bukan representasi kondisi lapangan. '
    'Kerangka NGBoost+DiCE bersifat transferable ke dataset air tabular apapun.'
)

doc.add_page_break()

# === REVISI 13 ===
doc.add_heading('REVISI 13: PERUBAHAN REFERENSI', level=2)
doc.add_paragraph('1. TAMBAH Patel et al. (2022) doi:10.1155/2022/9283293')
doc.add_paragraph('2. TAMBAH Yurtsever & Emec (2023)')
doc.add_paragraph('3. GANTI [5] Li soil moisture -> Duan et al. (2020)')
doc.add_paragraph('4. VERIFIKASI [2] Al Bataineh DOI')

doc.add_page_break()

# === REVISI 14 ===
doc.add_heading('REVISI 14: URUTAN PREPROCESSING (KRITIS)', level=2)
doc.add_paragraph('Lokasi: Sub-bab 3.3')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('URUTAN BENAR: ').bold = True
p.add_run('EDA -> Imputation(MICE) -> Split(70/15/15) -> Scaling(fit train) -> SMOTE(train only)')
doc.add_paragraph()
doc.add_paragraph('Sub-bab 3.3.3: "Split SETELAH imputation, SEBELUM scaling. Data validasi untuk tuning, data uji untuk evaluasi final."')
doc.add_paragraph('Sub-bab 3.3.4: "Parameter mu, sigma HANYA dari train. Transform val/test. Scaling SEBELUM SMOTE."')

doc.add_page_break()

# === REVISI 15 ===
doc.add_heading('REVISI 15: CONSTRAINT DOMAIN', level=2)
doc.add_paragraph('Lokasi: Sub-bab 3.5.2')
doc.add_paragraph()
doc.add_paragraph('Feasibility constraints berdasarkan Permenkes No. 2/2023 (primary) + WHO Guidelines 2022 (komplementer untuk Conductivity dan Organic Carbon).')
doc.add_paragraph()

headers15 = ['Parameter','Sumber','Range','Catatan']
rows15 = [
    ['pH','Permenkes','6.5-8.5','Langsung'],
    ['Hardness','Permenkes','0-500 mg/L','CaCO3'],
    ['TDS','Permenkes','0-1000 mg/L','Langsung'],
    ['Chloramines','Permenkes*','0-5 mg/L','*Dari Sisa Klor'],
    ['Sulfate','Permenkes','0-400 mg/L','Langsung'],
    ['Conductivity','WHO(2022)','0-1500 uS/cm','Tidak di Permenkes'],
    ['Organic Carbon','WHO(2022)','0-25 mg/L','Tidak langsung(KMnO4)'],
    ['Trihalomethanes','Permenkes','0-80 ug/L','Langsung'],
    ['Turbidity','Permenkes','0-5 NTU','Langsung'],
]
add_table(doc, headers15, rows15)

doc.add_paragraph('Catatan: Distribusi dataset melampaui batas regulasi (TDS hingga 61.227 vs max 1.000). Instance jauh dari range menghasilkan feasibility rendah - ini menjadi temuan penelitian.')

doc.add_page_break()

# === REVISI 16 ===
doc.add_heading('REVISI 16: BASELINE KOMPARATIF', level=2)
doc.add_paragraph('Lokasi: Sub-bab 3.4.5')
doc.add_paragraph()
doc.add_paragraph('Tambahkan: "Analisis komparatif 2 dimensi: (1) klasifikasi: Acc, Prec, Rec, F1; (2) probabilistik: NLL, ECE. Menjawab apakah NGBoost memberi nilai tambah vs deterministik."')

doc.add_page_break()

# === REVISI 17 ===
doc.add_heading('REVISI 17: KRITERIA DATA', level=2)
doc.add_paragraph('Lokasi: Sub-bab 3.2 (SISIPKAN 2 kalimat di dalam paragraf narasi dataset, BUKAN sub-bab baru)')
doc.add_paragraph('Alasan: Perlu justifikasi penggunaan seluruh data')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Sisipkan SETELAH kalimat tentang "3.276 baris sampel":').bold = True
doc.add_paragraph(
    'Seluruh record digunakan karena memiliki variabel target (Potability) yang valid dan tidak ditemukan record duplikat maupun record dengan seluruh fitur yang missing. '
    'Record dengan missing values parsial (pH: 14.98%, Sulfate: 23.84%, Trihalomethanes: 4.95%) tetap dipertahankan karena akan ditangani pada tahap imputasi MICE.'
)
doc.add_paragraph()
doc.add_paragraph('CATATAN: Ini BUKAN sub-bab baru. Hanya 2 kalimat yang disisipkan di paragraf yang sudah ada.')

doc.add_page_break()

# === REVISI 18 ===
doc.add_heading('REVISI 18: TABEL TOOLS', level=2)
doc.add_paragraph('Lokasi: Sub-bab 3.6 Tools dan Lingkungan Pengembangan')
doc.add_paragraph('Alasan: Narasi saat ini kurang terstruktur. GANTI narasi menjadi tabel.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('TEKS REVISI (ganti SELURUH sub-bab 3.6):').bold = True
doc.add_paragraph('Implementasi penelitian menggunakan lingkungan pengembangan berbasis Python pada Google Colab. Tabel 3.X menyajikan spesifikasi tools yang digunakan.')
doc.add_paragraph()

headers18 = ['Tools','Versi','Fungsi','Lisensi']
rows18 = [
    ['Python','3.9+','Bahasa pemrograman utama','PSF License'],
    ['Google Colab','Runtime','Lingkungan eksekusi (GPU)','Free/Pro'],
    ['ngboost','latest','Pemodelan probabilistik NGBoost','Apache 2.0'],
    ['dice-ml','latest','Pembangkitan counterfactual','MIT'],
    ['scikit-learn','latest','Preprocessing, baseline (XGB/RF), evaluasi','BSD-3'],
    ['pandas','latest','Manipulasi data tabular','BSD-3'],
    ['numpy','latest','Komputasi numerik','BSD-3'],
    ['matplotlib','latest','Visualisasi','PSF'],
    ['seaborn','latest','Visualisasi statistik','BSD-3'],
    ['imbalanced-learn','latest','SMOTE-ENN','MIT'],
]
add_table(doc, headers18, rows18)

doc.add_paragraph('Seluruh library berlisensi open-source sehingga tidak memerlukan izin khusus dari pengembang. Reproducibility dijamin melalui penggunaan random seed yang konsisten dan dokumentasi environment.')

doc.add_page_break()

# === REVISI 19 ===
doc.add_heading('REVISI 19: FIX DUPLIKASI TEORI BAB 3', level=2)
doc.add_paragraph('Lokasi: Sub-bab 3.4.1 dan 3.5.1')
doc.add_paragraph('Alasan: Sub-bab 3.4.1 "Landasan Teori Natural Gradient Boosting" MENDUPLIKASI BAB 2 (2.2.3). Sub-bab 3.5.1 "Landasan Teori" DiCE MENDUPLIKASI BAB 2 (2.2.4, 2.2.5).')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('A. Sub-bab 3.4.1 - UBAH JUDUL dan ISI:').bold = True
doc.add_paragraph('Judul lama: "3.4.1. Landasan Teori Natural Gradient Boosting"')
doc.add_paragraph('Judul baru: "3.4.1. Konfigurasi Model NGBoost"')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Isi baru (GANTI seluruhnya - fokus PROSEDUR bukan teori):').bold = True
doc.add_paragraph(
    'Konfigurasi NGBoost pada penelitian ini menggunakan distribusi Bernoulli sebagai output distribution untuk klasifikasi biner, dengan scoring rule berupa Negative Log-Likelihood. '
    'Base learner yang digunakan adalah decision tree dangkal (max_depth=3-5) untuk mencegah overfitting. '
    'Proses optimasi menggunakan Natural Gradient sebagaimana dijelaskan pada Subbab 2.2.3. '
    'Hyperparameter utama yang di-tuning meliputi: (1) n_estimators: jumlah iterasi boosting, (2) learning_rate: kecepatan konvergensi, (3) minibatch_frac: proporsi sampel per iterasi, (4) max_depth base learner. '
    'Tuning dilakukan menggunakan data validasi dengan kriteria NLL minimum dan ECE < 0.10. '
    'Early stopping diterapkan berdasarkan stabilisasi NLL pada data validasi.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('B. Sub-bab 3.5.1 - UBAH JUDUL dan ISI:').bold = True
doc.add_paragraph('Judul lama: "3.5.1. Landasan Teori"')
doc.add_paragraph('Judul baru: "3.5.1. Konfigurasi DiCE"')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Isi baru (GANTI seluruhnya - fokus PROSEDUR):').bold = True
doc.add_paragraph(
    'Framework DiCE dikonfigurasi dengan parameter berikut: (1) method: \'gradient\' untuk optimasi berbasis gradient pada model differentiable, '
    '(2) total_CFs: 4-5 counterfactual per instance, (3) desired_class: 1 (Potable/Layak), '
    '(4) features_to_vary: seluruh 9 fitur fisikokimia (semua dianggap actionable), '
    '(5) permitted_range: sesuai constraint domain pada Tabel 3.X (Permenkes + WHO). '
    'Fungsi loss DiCE menyeimbangkan validity, proximity, sparsity, diversity, dan feasibility sebagaimana dijelaskan pada Subbab 2.2.4 dan 2.2.5. '
    'Instance input adalah sampel dari data uji yang diprediksi sebagai Class 0 (Tidak Layak) oleh model NGBoost.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('PENJELASAN PERUBAHAN:').bold = True
doc.add_paragraph('- Teori matematika TETAP di BAB 2 (tidak diulang)')
doc.add_paragraph('- BAB 3 fokus pada PROSEDUR dan KONFIGURASI spesifik')
doc.add_paragraph('- Referensi silang ke BAB 2 ("sebagaimana dijelaskan pada Subbab 2.2.X") menghubungkan tanpa menduplikasi')

doc.add_page_break()

# === REVISI 20 ===
doc.add_heading('REVISI 20: URGENSI DENGAN DATA INDONESIA', level=2)
doc.add_paragraph('Lokasi: BAB 1, Paragraf PERTAMA Latar Belakang (halaman 1), SISIPKAN setelah kalimat pertama "Kualitas air minum yang memenuhi standar kelayakan merupakan aspek fundamental dalam menjaga kesehatan masyarakat sekaligus mendukung keberlanjutan ekosistem perairan."')
doc.add_paragraph('Alasan: Urgensi terlalu abstrak dan akademis. Tidak ada data konkret yang menunjukkan masalah nyata.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('KALIMAT YANG DISISIPKAN (tambahkan SETELAH kalimat pertama paragraf 1):').bold = True
doc.add_paragraph(
    'Menurut WHO (2022), sekitar 2 miliar orang di seluruh dunia masih menggunakan sumber air minum yang terkontaminasi. '
    'Di Indonesia, data Badan Pusat Statistik (2023) menunjukkan bahwa akses terhadap air minum layak belum merata di seluruh wilayah, dengan disparitas signifikan antara perkotaan dan perdesaan. '
    'Kondisi ini menegaskan kebutuhan akan sistem monitoring dan evaluasi kualitas air yang tidak hanya akurat dalam memprediksi kelayakan, tetapi juga mampu memberikan rekomendasi perbaikan yang spesifik dan dapat ditindaklanjuti oleh operator pengolahan air.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('CATATAN PENTING:').bold = True
p.add_run(' Verifikasi angka dan tahun dari sumber resmi:')
doc.add_paragraph('- WHO: "Drinking-water Key Facts" (2022/2023) - https://www.who.int/news-room/fact-sheets/detail/drinking-water')
doc.add_paragraph('- BPS: "Statistik Lingkungan Hidup Indonesia" atau "Indikator Perumahan dan Kesehatan Lingkungan" edisi terbaru')
doc.add_paragraph('- Pastikan sitasi sesuai format IEEE yang digunakan di proposal')

doc.add_page_break()

# === REVISI 21 ===
doc.add_heading('REVISI 21: JUSTIFIKASI VARIABEL FISIKOKIMIA (ACTIONABLE)', level=2)
doc.add_paragraph('Lokasi: BAB 1, Paragraf ke-2 Latar Belakang (halaman 1), SISIPKAN setelah kalimat "Parameter-parameter tersebut saling memengaruhi secara kompleks, sehingga status kelayakan air tidak dapat ditentukan melalui pendekatan sederhana dan memerlukan analisis data yang lebih mendalam."')
doc.add_paragraph('Alasan: Tidak ada justifikasi mengapa parameter fisikokimia dipilih. Dosen bisa bertanya "kenapa tidak parameter biologis?"')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('KALIMAT YANG DISISIPKAN:').bold = True
doc.add_paragraph(
    'Parameter fisikokimia dipilih sebagai fokus analisis karena bersifat controllable dan actionable, yaitu nilainya dapat diintervensi secara langsung melalui proses pengolahan air seperti koagulasi, filtrasi, aerasi, dan disinfeksi [1], [6]. '
    'Karakteristik ini menjadikan parameter fisikokimia sesuai dengan kebutuhan analisis preskriptif berbasis counterfactual, yang memerlukan fitur-fitur yang secara operasional dapat dimodifikasi untuk mencapai kondisi target yang diinginkan.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('PENJELASAN:').bold = True
p.add_run(' Kalimat ini memberikan justifikasi GANDA:')
doc.add_paragraph('(1) Fisikokimia bersifat controllable (bisa diubah di lapangan)')
doc.add_paragraph('(2) Counterfactual memerlukan fitur actionable (syarat metodologis)')
doc.add_paragraph('Sehingga pilihan variabel BUKAN kebetulan dari dataset, tapi karena memang sesuai dengan kebutuhan metodologi penelitian.')

doc.add_page_break()

# === REVISI 22 ===
doc.add_heading('REVISI 22: TRANSISI EKSPLISIT KE RUMUSAN MASALAH', level=2)
doc.add_paragraph('Lokasi: BAB 1, Sub-bab 1.2 Rumusan Masalah (halaman 3), GANTI paragraf pengantar sebelum daftar rumusan masalah.')
doc.add_paragraph('Alasan: Transisi dari latar belakang ke rumusan masalah kurang eksplisit. Derivasi RM dari gap tidak ditunjukkan secara jelas.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('PARAGRAF LAMA:').bold = True
doc.add_paragraph(
    '"Berdasarkan latar belakang yang telah diuraikan, ketidakpastian pada data parameter fisikokimia, karakteristik black-box model prediktif, serta ketiadaan mekanisme rekomendasi preskriptif yang dapat ditindaklanjuti menjadi permasalahan utama dalam penelitian ini. Oleh karena itu, rumusan masalahnya adalah sebagai berikut:"'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('PARAGRAF BARU (GANTI):').bold = True
doc.add_paragraph(
    '"Berdasarkan latar belakang yang telah diuraikan, teridentifikasi tiga kesenjangan utama yang menjadi fokus penelitian ini: '
    '(1) ketiadaan estimasi ketidakpastian pada model machine learning deterministik yang digunakan di domain kualitas air, '
    '(2) ketiadaan mekanisme preskriptif yang mampu mentransformasi output prediksi menjadi rekomendasi tindakan operasional untuk perbaikan kualitas air, dan '
    '(3) belum adanya implementasi counterfactual explanations pada data fisikokimia air meskipun potensinya telah diidentifikasi dalam systematic review [7]. '
    'Berdasarkan tiga kesenjangan tersebut, dirumuskan pertanyaan penelitian sebagai berikut:"'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('PENJELASAN PERUBAHAN:').bold = True
doc.add_paragraph('- Menghilangkan kata "black-box" (sesuai revisi sebelumnya)')
doc.add_paragraph('- Menunjukkan 3 gap secara eksplisit dan terstruktur (numbered)')
doc.add_paragraph('- Setiap gap langsung terhubung ke 1 rumusan masalah')
doc.add_paragraph('- Menambahkan referensi [7] (Aderemi systematic review) sebagai backing')

doc.add_page_break()

# === REVISI 23 (BARU) ===
doc.add_heading('REVISI 23: TABEL PERBANDINGAN XAI', level=2)
doc.add_paragraph('Lokasi: Sub-bab 2.2.4 (sisipkan sebagai tabel perbandingan)')
doc.add_paragraph('Alasan: Perlu justifikasi akademis mengapa DiCE dipilih dibandingkan metode XAI lain.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Tabel Perbandingan SHAP vs LIME vs Counterfactual (DiCE):').bold = True
doc.add_paragraph()

headers23 = ['Metode','Tipe','Output','Kelebihan','Keterbatasan','Relevansi']
rows23 = [
    ['SHAP','Post-hoc diagnostik','Feature importance','Global+local explanation','Tidak preskriptif, hanya menjelaskan WHY','Tidak sesuai'],
    ['LIME','Post-hoc diagnostik','Linear approximation','Model-agnostic','Tidak stabil, tidak preskriptif','Tidak sesuai'],
    ['CF (DiCE)','Preskriptif aktif','Perubahan fitur spesifik','Actionable, diverse, constrained','Bergantung kualitas model','Sesuai'],
]
add_table(doc, headers23, rows23)

p = doc.add_paragraph()
p.add_run('Kesimpulan: ').bold = True
p.add_run('DiCE dipilih karena merupakan satu-satunya metode yang bersifat preskriptif (memberikan HOW, bukan hanya WHY). SHAP dan LIME hanya bersifat diagnostik - menjelaskan fitur mana yang berpengaruh, namun tidak memberikan rekomendasi perubahan spesifik yang actionable untuk memperbaiki kualitas air.')

doc.add_page_break()

# === REVISI 24 (BARU) ===
doc.add_heading('REVISI 24: JUSTIFIKASI NGBoost', level=2)
doc.add_paragraph('Lokasi: Akhir sub-bab 2.2.3 (tambahkan paragraf justifikasi)')
doc.add_paragraph('Alasan: Dosen kemungkinan bertanya "Kenapa NGBoost? Kenapa bukan BNN atau deep learning?"')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Paragraf yang ditambahkan:').bold = True
doc.add_paragraph(
    'NGBoost dipilih sebagai model probabilistik dalam penelitian ini berdasarkan pertimbangan berikut: '
    '(1) Bayesian Neural Network (BNN) memiliki kompleksitas komputasi tinggi dan sulit di-tune untuk dataset tabular berukuran kecil-sedang; '
    '(2) MC Dropout hanya memberikan approximate uncertainty dan memerlukan arsitektur neural network yang tidak optimal untuk data tabular; '
    '(3) Deep learning secara umum menunjukkan performa inferior dibandingkan gradient boosting pada dataset tabular berukuran kecil (Grinsztajn et al., 2022); '
    '(4) NGBoost secara native menghasilkan distribusi probabilitas penuh melalui Natural Gradient descent, bukan hanya point estimate; '
    '(5) NGBoost kompatibel dengan framework DiCE melalui method predict_proba yang diperlukan untuk pembangkitan counterfactual. '
    'Kombinasi performa kompetitif pada data tabular dan kompatibilitas dengan DiCE menjadikan NGBoost pilihan yang optimal untuk kerangka preskriptif ini.'
)

doc.add_page_break()

# === REVISI 25 (BARU) ===
doc.add_heading('REVISI 25: JUSTIFIKASI DiCE', level=2)
doc.add_paragraph('Lokasi: Akhir sub-bab 2.2.5 (tambahkan paragraf justifikasi)')
doc.add_paragraph('Alasan: Perlu justifikasi mengapa DiCE dipilih dibandingkan metode counterfactual lain.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Paragraf yang ditambahkan:').bold = True
doc.add_paragraph(
    'DiCE dipilih sebagai framework counterfactual dalam penelitian ini berdasarkan pertimbangan berikut: '
    '(1) Metode Wachter et al. (2017) hanya menghasilkan 1 counterfactual per instance, tidak memberikan alternatif rekomendasi; '
    '(2) FACE (Feasible Actionable Counterfactual Explanations) kurang fleksibel dalam mendefinisikan constraints domain-specific; '
    '(3) Growing Spheres tidak mendukung multi-objective optimization (validity + proximity + diversity secara simultan); '
    '(4) DiCE menghasilkan diverse counterfactuals (k>1) yang memberikan multiple actionable paths; '
    '(5) DiCE mendukung permitted_range untuk constraint domain (Permenkes/WHO) dan features_to_vary untuk kontrol actionability; '
    '(6) Library dice-ml sudah mature, well-documented, dan compatible dengan scikit-learn API yang digunakan NGBoost. '
    'Keunggulan DiCE dalam menghasilkan diverse, constrained, dan actionable counterfactuals menjadikannya pilihan optimal untuk domain kualitas air.'
)

doc.add_page_break()

# === REVISI 26 (BARU) ===
doc.add_heading('REVISI 26: BATASAN KONSEP PROBABILISTIK', level=2)
doc.add_paragraph('Lokasi: Sub-bab 2.2.3 (sisipkan di awal atau setelah penjelasan NGBoost)')
doc.add_paragraph('Alasan: Menghindari miskonsepsi bahwa "probabilistik" = Bayesian inference penuh.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Teks yang disisipkan:').bold = True
doc.add_paragraph(
    'Catatan penting: "Probabilistik" dalam konteks penelitian ini merujuk pada pemodelan distribusi Bernoulli melalui NGBoost, '
    'BUKAN Bayesian inference penuh (seperti BNN, Gaussian Process, atau MC Dropout). '
    'Uncertainty yang diestimasi bersifat aleatoric (ketidakpastian data), yang direpresentasikan sebagai mu*(1-mu) dari distribusi Bernoulli, '
    'bukan epistemic uncertainty (ketidakpastian model). '
    'Distingsi ini penting agar tidak terjadi overclaim terhadap kemampuan uncertainty estimation dari model yang digunakan.'
)

doc.add_page_break()

# === REVISI 27 (BARU) ===
doc.add_heading('REVISI 27: WHAT-IF ANALYSIS', level=2)
doc.add_paragraph('Lokasi: Sub-bab 2.3 (sisipkan sebelum paragraf penutup)')
doc.add_paragraph('Alasan: Memperkuat argumentasi mengapa integrasi NGBoost+DiCE diperlukan vs alternatif.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Teks yang disisipkan:').bold = True
doc.add_paragraph(
    'Untuk memperjelas kontribusi integrasi yang diusulkan, berikut analisis what-if terhadap alternatif: '
    '(1) Jika menggunakan XGBoost saja: model menghasilkan prediksi deterministik tanpa estimasi uncertainty, sehingga operator tidak mengetahui tingkat keyakinan prediksi; '
    '(2) Jika menggunakan SHAP saja: output hanya bersifat diagnostik (menjelaskan fitur mana yang berpengaruh), tanpa memberikan rekomendasi perubahan spesifik yang actionable; '
    '(3) Integrasi NGBoost + DiCE mengatasi kedua limitasi tersebut secara simultan - NGBoost memberikan prediksi dengan uncertainty estimation, dan DiCE mentransformasi prediksi tersebut menjadi rekomendasi preskriptif yang actionable dan constrained.'
)

doc.add_page_break()

# === REVISI 28 (BARU) ===
doc.add_heading('REVISI 28: NOVELITAS STATEMENT', level=2)
doc.add_paragraph('Lokasi: Kalimat terakhir sub-bab 2.3')
doc.add_paragraph('Alasan: Perlu kalimat penutup yang menegaskan novelitas/kontribusi utama penelitian.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Ganti kalimat penutup sub-bab 2.3 menjadi:').bold = True
doc.add_paragraph(
    '"Penelitian ini merupakan first practical implementation yang mengintegrasikan NGBoost sebagai model probabilistik dengan DiCE sebagai framework preskriptif pada domain kualitas air, '
    'dengan feasibility constraints berbasis regulasi nasional (Permenkes No. 2/2023) dan internasional (WHO Guidelines 2022). '
    'Kontribusi ini menjembatani kesenjangan antara kemampuan prediktif machine learning dan kebutuhan operasional di lapangan."'
)

doc.add_page_break()

# === STRATEGI SIDANG ===
doc.add_heading('STRATEGI SIDANG', level=2)

doc.add_heading('Pertanyaan Dataset', level=3)
doc.add_paragraph('L1: "Dataset publik Kaggle, CC0, 99+ paper Scopus termasuk Patel et al. (2022)."')
doc.add_paragraph('L2: "Benchmark. Tidak klaim representatifitas. Fokus metodologi."')
doc.add_paragraph('L3: "Kerangka NGBoost+DiCE transferable ke data PDAM/lab."')

doc.add_heading('Pertanyaan Constraint Regulasi', level=3)
doc.add_paragraph('L1: "Permenkes No. 2/2023 + WHO untuk 2 parameter yang tidak ter-cover."')
doc.add_paragraph('L2: "Conductivity dan Organic Carbon tidak di Permenkes, gunakan WHO komplementer."')
doc.add_paragraph('L3: "Distribusi dataset melampaui regulasi = temuan bahwa kerangka lebih cocok untuk data real."')

doc.add_heading('Pertanyaan Populasi/Sampel/Instrumen', level=3)
doc.add_paragraph('L1: "Data sekunder 3.276 record. Populasi = seluruh dataset benchmark."')
doc.add_paragraph('L2: "Bukan penelitian survei - instrumen = library open-source (ngboost, dice-ml, scikit-learn). Lisensi Apache/MIT/BSD, tidak perlu izin."')
doc.add_paragraph('L3: "Reliabilitas = reproducibility: random seed fixed, environment terdokumentasi, code publishable."')

doc.add_heading('Pertanyaan Data Leakage', level=3)
doc.add_paragraph('L1: "Urutan: Imputation -> Split -> Scaling -> SMOTE. Scaling fit hanya pada train."')
doc.add_paragraph('L2: "SMOTE hanya pada train set. Data validasi dan uji tidak tersentuh sampai evaluasi."')
doc.add_paragraph('L3: "Imputation sebelum split karena MICE bersifat unsupervised (tidak menggunakan label target)."')

doc.add_heading('Pertanyaan Variabel & Topik', level=3)
doc.add_paragraph('L1: "Parameter fisikokimia bersifat controllable/actionable - bisa diintervensi melalui proses pengolahan air."')
doc.add_paragraph('L2: "Counterfactual explanations memerlukan fitur yang bisa diubah nilainya. Parameter biologis/mikroba tidak actionable karena bersifat outcome, bukan input yang bisa di-treat langsung."')
doc.add_paragraph('L3: "Pemilihan topik berdasarkan gap yang eksplisit disebutkan di systematic review Aderemi et al. (2025) - belum ada implementasi CF di domain air."')

doc.add_heading('Pertanyaan Teori & Metode', level=3)
doc.add_paragraph('L1: "NGBoost = probabilistik via Bernoulli + Natural Gradient. DiCE = preskriptif via diverse counterfactuals dengan constraints. Integrasi keduanya = prediksi + uncertainty + rekomendasi actionable."')
doc.add_paragraph('L2: "NGBoost dipilih vs BNN (terlalu kompleks untuk tabular kecil), vs MC Dropout (approximate), vs deep learning (inferior pada tabular - Grinsztajn 2022). DiCE dipilih vs Wachter (hanya 1 CF), vs FACE (kurang flexible constraints), vs Growing Spheres (tidak multi-objective)."')
doc.add_paragraph('L3: "Probabilistik di sini = aleatoric uncertainty via Bernoulli mu*(1-mu), BUKAN epistemic/Bayesian. What-if: XGBoost saja = tanpa uncertainty, SHAP saja = diagnostik only. NGBoost+DiCE = solusi simultan untuk kedua limitasi. First implementation di domain air dengan constraints Permenkes+WHO."')

doc.add_page_break()

# === CHECKLIST ===
doc.add_heading('CHECKLIST REVISI (28 Item)', level=2)

headers_cl = ['No','Bagian','Lokasi','Jenis','Prioritas','Status']
rows_cl = [
    ['1','Abstrak+Hipotesa','Hal.i','Rewrite','TINGGI','[ ]'],
    ['2','Latar Belakang','Hal.1','Ganti paragraf','TINGGI','[ ]'],
    ['3','Rumusan Masalah','Sub.1.2','Rewrite','TINGGI','[ ]'],
    ['4','Tujuan','Sub.1.3','Rewrite','TINGGI','[ ]'],
    ['5','Hapus Batasan','Sub.1.4','Hapus+renomor','SEDANG','[ ]'],
    ['6','Tabel Penelitian','Tabel 2.1','Ganti tabel','TINGGI','[ ]'],
    ['7','Sitasi Gambar','Gambar 2.1','Tambah sitasi','SEDANG','[ ]'],
    ['8','Rumus Metrik','Sub.2.2.6','Rewrite+rumus','TINGGI','[ ]'],
    ['9','Flowchart','Gambar 3.1','Gambar ulang(27 node)','TINGGI','[ ]'],
    ['10','Tabel Preprocessing','Sub.3.3','Tambah 3 tabel','TINGGI','[ ]'],
    ['11','Penomoran BAB 3','Sub.3.5-3.6','Fix struktur','SEDANG','[ ]'],
    ['12','Narasi Dataset','Sub.3.2','Ubah narasi','TINGGI','[ ]'],
    ['13','Referensi','Daftar Pustaka','Tambah/ganti','TINGGI','[ ]'],
    ['14','Urutan Preprocessing','Sub.3.3','Fix urutan','KRITIS','[ ]'],
    ['15','Constraint Domain','Sub.3.5.2','Permenkes+WHO','TINGGI','[ ]'],
    ['16','Baseline Komparatif','Sub.3.4.5','Tambah analisis','SEDANG','[ ]'],
    ['17','Kriteria Data','Sub.3.2','Sisipkan 2 kalimat','SEDANG','[ ]'],
    ['18','Tabel Tools','Sub.3.6','Replace narasi->tabel','SEDANG','[ ]'],
    ['19','Fix Duplikasi Teori','Sub.3.4.1+3.5.1','Ubah judul+isi','TINGGI','[ ]'],
    ['20','Urgensi Indonesia','Hal.1 par.1','Sisipkan data WHO/BPS','SEDANG','[ ]'],
    ['21','Justifikasi Variabel','Hal.1 par.2','Sisipkan 2 kalimat','SEDANG','[ ]'],
    ['22','Transisi ke RM','Hal.3 sub.1.2','Ganti paragraf pengantar','SEDANG','[ ]'],
    ['23','Tabel Perbandingan XAI','Sub.2.2.4','Tambah tabel SHAP vs LIME vs DiCE','TINGGI','[ ]'],
    ['24','Justifikasi NGBoost','Sub.2.2.3','Tambah paragraf justifikasi','TINGGI','[ ]'],
    ['25','Justifikasi DiCE','Sub.2.2.5','Tambah paragraf justifikasi','TINGGI','[ ]'],
    ['26','Batasan Probabilistik','Sub.2.2.3','Sisipkan klarifikasi konsep','SEDANG','[ ]'],
    ['27','What-If Analysis','Sub.2.3','Sisipkan analisis alternatif','SEDANG','[ ]'],
    ['28','Novelitas Statement','Sub.2.3','Ganti kalimat penutup','TINGGI','[ ]'],
]
add_table(doc, headers_cl, rows_cl)

doc.add_page_break()

# === CATATAN PENTING ===
doc.add_heading('CATATAN PENTING (12 item)', level=2)
doc.add_paragraph('1. Verifikasi metrik dari paper asli')
doc.add_paragraph('2. Tabel preprocessing ilustratif - ganti data aktual saat implementasi')
doc.add_paragraph('3. Flowchart digambar manual (draw.io)')
doc.add_paragraph('4. Semua gambar perlu sitasi')
doc.add_paragraph('5. Referensi baru verifikasi DOI')
doc.add_paragraph('6. JANGAN klaim dataset sebagai data lapangan')
doc.add_paragraph('7. URUTAN KRITIS: Imputation->Split->Scaling->SMOTE')
doc.add_paragraph('8. Permenkes No. 2/2023 - verifikasi dari sumber resmi')
doc.add_paragraph('9. BAB 3 = PROSEDUR, BAB 2 = TEORI. Jangan duplikasi.')
doc.add_paragraph('10. Data WHO dan BPS HARUS diverifikasi dari sumber resmi sebelum dimasukkan ke proposal')
doc.add_paragraph('11. Manfaat penelitian TIDAK perlu ditambahkan sebagai sub-bab (tidak ada di template) - siapkan jawaban lisan saja')
doc.add_paragraph('12. Justifikasi metode (NGBoost, DiCE) harus EKSPLISIT di BAB 2 - jangan hanya deskripsi teori tanpa alasan pemilihan. Siapkan jawaban 3 level untuk pertanyaan "kenapa metode ini?"')

# Save
doc.save('/projects/sandbox/ProposalSeminarS6/REVISI_PROPOSAL_AFLAH.docx')
print("DONE: REVISI_PROPOSAL_AFLAH.docx v6 created successfully.")
