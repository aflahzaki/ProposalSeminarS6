"""
Script to complete the proposal V8 docx:
1. Fill Tabel 1.1 (Jadwal Kegiatan) with Gantt chart shading
2. Add reference placeholders [ref Mothilal2020], [ref Canada-Figshare], [ref Guo2017]
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FILENAME = 'AflahZakISiregar_103062300095_PenulisanProposal-Final-AI-V8.docx'


def set_cell_shading(cell, color="4472C4"):
    """Apply shading/fill color to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    tcPr.append(shading)


def fill_gantt_chart(doc):
    """Fill Tabel 1.1 (doc.tables[1]) with Gantt chart shading."""
    table = doc.tables[1]
    
    # Gantt schedule: (row_index, start_col, end_col_inclusive)
    # Columns: 2-5 = Month 1 (4 weeks), 6-9 = Month 2, 10-13 = Month 3,
    #           14-17 = Month 4, 18-21 = Month 5, 22-25 = Month 6
    gantt_schedule = {
        2: (2, 9),    # Studi Literatur: months 1-2 (cols 2-9)
        3: (6, 13),   # Pengumpulan dan Pengolahan Data: months 2-3 (cols 6-13)
        4: (10, 17),  # Pemodelan NGBoost: months 3-4 (cols 10-17)
        5: (14, 21),  # Implementasi DiCE: months 4-5 (cols 14-21)
        6: (18, 21),  # Pengujian dan Evaluasi: month 5 (cols 18-21)
        7: (14, 25),  # Penyusunan Laporan: months 4-6 (cols 14-25)
    }
    
    color = "4472C4"  # Microsoft Office blue
    
    for row_idx, (start_col, end_col) in gantt_schedule.items():
        row = table.rows[row_idx]
        for col_idx in range(start_col, end_col + 1):
            cell = row.cells[col_idx]
            set_cell_shading(cell, color)
    
    print(f"Gantt chart filled for rows 2-7 with blue shading")


def add_reference_placeholders(doc):
    """Add reference placeholders to relevant paragraphs."""
    
    # [ref Canada-Figshare] - Paragraph 211 about Canada dataset
    para_211 = doc.paragraphs[211]
    target_text = "repositori Nature.com (Scientific Data)"
    if target_text in para_211.text and "[ref Canada-Figshare]" not in para_211.text:
        # Find the run containing the target text and append reference
        for run in para_211.runs:
            if "Scientific Data)" in run.text:
                run.text = run.text.replace(
                    "Scientific Data)",
                    "Scientific Data) [ref Canada-Figshare]"
                )
                print("Added [ref Canada-Figshare] to paragraph 211")
                break
        else:
            # If we can't find the exact run, append to the last run
            para_211.runs[-1].text += " [ref Canada-Figshare]"
            print("Added [ref Canada-Figshare] to paragraph 211 (appended)")
    elif "[ref Canada-Figshare]" in para_211.text:
        print("[ref Canada-Figshare] already present in paragraph 211")
    else:
        print(f"WARNING: Could not find target text in paragraph 211")
    
    # [ref Guo2017] - Paragraph 237 about ECE
    para_237 = doc.paragraphs[237]
    target_text_237 = "stabilisasi ECE"
    if target_text_237 in para_237.text and "[ref Guo2017]" not in para_237.text:
        for run in para_237.runs:
            if "ECE" in run.text and "stabilisasi" in run.text:
                run.text = run.text.replace(
                    "stabilisasi ECE",
                    "stabilisasi ECE [ref Guo2017]"
                )
                print("Added [ref Guo2017] to paragraph 237")
                break
        else:
            # Try finding ECE in any run
            for run in para_237.runs:
                if "ECE" in run.text:
                    run.text = run.text.replace("ECE", "ECE [ref Guo2017]", 1)
                    print("Added [ref Guo2017] to paragraph 237 (at ECE mention)")
                    break
            else:
                para_237.runs[-1].text += " [ref Guo2017]"
                print("Added [ref Guo2017] to paragraph 237 (appended)")
    elif "[ref Guo2017]" in para_237.text:
        print("[ref Guo2017] already present in paragraph 237")
    else:
        print(f"WARNING: Could not find ECE in paragraph 237")
    
    # [ref Mothilal2020] - Paragraph 247 about DiCE
    para_247 = doc.paragraphs[247]
    target_text_247 = "Diverse Counterfactual explanations (DiCE)"
    if target_text_247 in para_247.text and "[ref Mothilal2020]" not in para_247.text:
        for run in para_247.runs:
            if "(DiCE)" in run.text:
                run.text = run.text.replace(
                    "(DiCE)",
                    "(DiCE) [ref Mothilal2020]"
                )
                print("Added [ref Mothilal2020] to paragraph 247")
                break
        else:
            # Try finding DiCE in any run
            for run in para_247.runs:
                if "DiCE" in run.text:
                    run.text = run.text.replace("DiCE", "DiCE [ref Mothilal2020]", 1)
                    print("Added [ref Mothilal2020] to paragraph 247 (at DiCE mention)")
                    break
            else:
                para_247.runs[-1].text += " [ref Mothilal2020]"
                print("Added [ref Mothilal2020] to paragraph 247 (appended)")
    elif "[ref Mothilal2020]" in para_247.text:
        print("[ref Mothilal2020] already present in paragraph 247")
    else:
        # Check if the target text is split across runs
        full_text = para_247.text
        if "DiCE" in full_text:
            for run in para_247.runs:
                if "DiCE" in run.text:
                    run.text = run.text.replace("DiCE", "DiCE [ref Mothilal2020]", 1)
                    print("Added [ref Mothilal2020] to paragraph 247 (at DiCE)")
                    break
        else:
            print(f"WARNING: Could not find DiCE in paragraph 247")


def main():
    doc = Document(FILENAME)
    
    print("=" * 60)
    print("Filling Gantt chart (Tabel 1.1)...")
    print("=" * 60)
    fill_gantt_chart(doc)
    
    print()
    print("=" * 60)
    print("Adding reference placeholders...")
    print("=" * 60)
    add_reference_placeholders(doc)
    
    print()
    print("=" * 60)
    print("Saving document...")
    print("=" * 60)
    doc.save(FILENAME)
    print(f"Document saved: {FILENAME}")
    
    # Verify
    print()
    print("=" * 60)
    print("Verification...")
    print("=" * 60)
    doc2 = Document(FILENAME)
    print(f"Document loads OK: {len(doc2.paragraphs)} paragraphs, {len(doc2.tables)} tables")
    
    # Check Gantt
    t = doc2.tables[1]
    shaded_count = 0
    for row_idx in range(2, 8):
        row = t.rows[row_idx]
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None and shd.get(qn('w:fill')) == '4472C4':
                    shaded_count += 1
    print(f"Gantt chart shaded cells: {shaded_count}")
    
    # Check references
    full_text = ' '.join(p.text for p in doc2.paragraphs)
    refs = ['[ref Mothilal2020]', '[ref Canada-Figshare]', '[ref Guo2017]']
    for ref in refs:
        if ref in full_text:
            print(f"  {ref} - FOUND")
        else:
            print(f"  {ref} - MISSING!")


if __name__ == '__main__':
    main()
